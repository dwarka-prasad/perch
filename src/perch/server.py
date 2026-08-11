"""Perch backend — HTTP server, system collectors, and dev tooling.

Serves the web UI (perch/web) and a JSON API on 127.0.0.1. Run via
``python -m perch`` or the ``perch`` console script.
"""

import base64
import fcntl
import glob
import hashlib
import json
import os
import pty as _pty
import pwd
import re
import secrets
import select
import shutil
import socket
import stat
import struct
import subprocess
import sys
import termios
import threading
import time
import urllib.error
import urllib.parse
from collections import deque
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

try:
    import psutil
except ImportError:  # self-install the only external dependency
    print("psutil missing — installing with pip --user …")
    subprocess.run([sys.executable, "-m", "pip", "install",
                    "--user", "psutil"], check=True)
    import psutil

PORT = int(os.environ.get("PERCH_PORT") or
           (sys.argv[1] if len(sys.argv) > 1 and sys.argv[1].isdigit()
            else 9080))
HOST = os.environ.get("PERCH_HOST", "127.0.0.1")
HOME = os.path.expanduser("~")
TOKEN_FILE = os.path.join(HOME, ".perch-token")

_SUBPROC_TIMEOUT = 15
_sp_run = subprocess.run


def _run(cmd, **kw):
    """subprocess.run with a default timeout so no request can hang."""
    kw.setdefault("timeout", _SUBPROC_TIMEOUT)
    try:
        return _sp_run(cmd, **kw)
    except subprocess.TimeoutExpired:
        return subprocess.CompletedProcess(cmd, 124, "", "command timed out")


def atomic_write(path, text, mode=0o600):
    """Write text to path via a temp file + rename, so a crash mid-write
    never leaves a truncated/corrupt file behind."""
    tmp = path + ".tmp"
    with open(tmp, "w") as f:
        f.write(text)
    os.chmod(tmp, mode)
    os.replace(tmp, path)


def _sd_notify(msg):
    """Tell systemd we're alive (Type=notify + WatchdogSec)."""
    addr = os.environ.get("NOTIFY_SOCKET")
    if not addr:
        return
    if addr.startswith("@"):
        addr = "\0" + addr[1:]
    try:
        s = socket.socket(socket.AF_UNIX, socket.SOCK_DGRAM)
        s.sendto(msg.encode(), addr)
        s.close()
    except OSError:
        pass


# failed-auth lockout: 20 bad tokens in 10 min locks that address out
_AUTH_FAILS = {}


def _auth_fail(ip):
    n, t0 = _AUTH_FAILS.get(ip, (0, time.time()))
    if time.time() - t0 > 600:
        n, t0 = 0, time.time()
    _AUTH_FAILS[ip] = (n + 1, t0)


def _auth_locked(ip):
    n, t0 = _AUTH_FAILS.get(ip, (0, 0))
    return n >= 20 and time.time() - t0 < 600
BOOT = psutil.boot_time()


def load_token():
    try:
        with open(TOKEN_FILE) as f:
            t = f.read().strip()
            if len(t) >= 16:
                return t
    except OSError:
        pass
    t = secrets.token_urlsafe(24)
    with open(TOKEN_FILE, "w") as f:
        f.write(t)
    os.chmod(TOKEN_FILE, 0o600)
    return t


TOKEN = load_token()

WEB_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "web")
STATIC_TYPES = {".css": "text/css", ".js": "text/javascript",
                ".svg": "image/svg+xml", ".png": "image/png",
                ".ico": "image/x-icon"}


def render_index(token):
    with open(os.path.join(WEB_DIR, "index.html"), encoding="utf-8") as f:
        return f.read().replace("{{TOKEN}}", token).replace("{{HOME}}", HOME)


# -------------------------------------------------------------------- gpu ----

GPU_CARD = next((os.path.dirname(p) for p in
                 glob.glob("/sys/class/drm/card*/gt_cur_freq_mhz")), None)


def _gpu_name():
    try:
        r = _run(["lspci"], capture_output=True, text=True, timeout=5)
        for line in r.stdout.splitlines():
            if "VGA" in line or "3D" in line:
                name = line.split(": ", 1)[1]
                return re.sub(r"\s*\(rev ..\)$", "", name)
    except Exception:  # noqa: BLE001
        pass
    return "GPU"


GPU_NAME = _gpu_name()
if "46a8" in GPU_NAME:
    GPU_NAME = "Intel Iris Xe Graphics (Alder Lake-P)"


def _gpu_freq(which):
    try:
        with open(os.path.join(GPU_CARD, f"gt_{which}_freq_mhz")) as f:
            return int(f.read())
    except (OSError, TypeError, ValueError):
        return None


# Per-process GPU busyness from /proc/<pid>/fdinfo DRM counters (own procs).
_gpu_pids = {}     # pid -> {"fds": [fd,...], "name": str}
_gpu_prev = {}     # pid -> (t, busy_ns)
_gpu_rates = {}    # pid -> {"name":, "busy": %}
_gpu_scan_at = 0.0


def _gpu_rescan_clients():
    global _gpu_pids, _gpu_scan_at
    _gpu_scan_at = time.time()
    found = {}
    me = os.getuid()
    for p in psutil.process_iter(["pid", "name", "uids"]):
        try:
            if not p.info["uids"] or p.info["uids"].real != me:
                continue
            pid = p.info["pid"]
            fddir = f"/proc/{pid}/fd"
            fds = []
            for fd in os.listdir(fddir):
                try:
                    if "/dev/dri/" in os.readlink(os.path.join(fddir, fd)):
                        fds.append(fd)
                except OSError:
                    continue
            if fds:
                found[pid] = {"fds": fds, "name": p.info["name"]}
        except (psutil.Error, OSError):
            continue
    _gpu_pids = found


def _gpu_pid_busy_ns(pid, fds):
    per_client = {}
    for fd in fds:
        try:
            with open(f"/proc/{pid}/fdinfo/{fd}") as f:
                cid, busy = None, None
                for line in f:
                    if line.startswith("drm-client-id:"):
                        cid = line.split(":")[1].strip()
                    elif line.startswith("drm-engine-render:"):
                        busy = int(line.split(":")[1].strip().split()[0])
                if cid is not None and busy is not None:
                    per_client[cid] = busy
        except (OSError, ValueError):
            continue
    return sum(per_client.values())


def _gpu_sample():
    """Returns total render busy %, updates per-process rates."""
    global _gpu_rates
    now = time.time()
    if now - _gpu_scan_at > 20:
        _gpu_rescan_clients()
    total_pct = 0.0
    rates = {}
    for pid, info in list(_gpu_pids.items()):
        busy = _gpu_pid_busy_ns(pid, info["fds"])
        prev = _gpu_prev.get(pid)
        _gpu_prev[pid] = (now, busy)
        if prev and busy >= prev[1] and now > prev[0]:
            pct = (busy - prev[1]) / ((now - prev[0]) * 1e9) * 100
            if pct > 0.05:
                rates[pid] = {"name": info["name"], "busy": round(pct, 1)}
            total_pct += pct
    for pid in list(_gpu_prev):
        if pid not in _gpu_pids:
            del _gpu_prev[pid]
    _gpu_rates = rates
    return min(100.0, total_pct)


def gpu_info():
    top = sorted(({"pid": k, **v} for k, v in _gpu_rates.items()),
                 key=lambda x: -x["busy"])[:8]
    status = None
    try:
        with open(os.path.join(GPU_CARD, "device/power/runtime_status")) as f:
            status = f.read().strip()
    except (OSError, TypeError):
        pass
    last = HISTORY[-1] if HISTORY else {}
    return {"name": GPU_NAME, "cur": _gpu_freq("cur"), "max": _gpu_freq("max"),
            "min": _gpu_freq("min"), "busy": last.get("gpu"),
            "status": status, "top": top}


# ---------------------------------------------------------------- sampler ----

HISTORY = deque(maxlen=150)  # ~5 min at 2 s
_prev_net = None


_prev_disk = None


def _max_temp():
    try:
        best = None
        for entries in (psutil.sensors_temperatures() or {}).values():
            for e in entries:
                if e.current and (best is None or e.current > best):
                    best = e.current
        return round(best, 1) if best else None
    except Exception:  # noqa: BLE001
        return None


def sampler():
    global _prev_net, _prev_disk
    psutil.cpu_percent(interval=None)  # prime
    while True:
        time.sleep(2)
        cpu = psutil.cpu_percent(interval=None)
        mem = psutil.virtual_memory().percent
        io = psutil.net_io_counters()
        now = time.time()
        if _prev_net:
            dt = now - _prev_net[0]
            down = max(0, io.bytes_recv - _prev_net[1]) / dt
            up = max(0, io.bytes_sent - _prev_net[2]) / dt
        else:
            down = up = 0.0
        _prev_net = (now, io.bytes_recv, io.bytes_sent)
        dio = psutil.disk_io_counters()
        if dio and _prev_disk:
            dt2 = now - _prev_disk[0]
            dread = max(0, dio.read_bytes - _prev_disk[1]) / dt2
            dwrite = max(0, dio.write_bytes - _prev_disk[2]) / dt2
        else:
            dread = dwrite = 0.0
        if dio:
            _prev_disk = (now, dio.read_bytes, dio.write_bytes)
        try:
            gpu = _gpu_sample() if GPU_CARD else None
        except Exception:  # noqa: BLE001
            gpu = None
        entry = {"t": now, "cpu": cpu, "mem": mem, "down": down,
                 "up": up, "gpu": gpu,
                 "gfreq": _gpu_freq("cur") if GPU_CARD else None,
                 "dr": dread, "dw": dwrite, "temp": _max_temp()}
        HISTORY.append(entry)
        try:
            monitor_tick(entry)
        except Exception:  # noqa: BLE001
            pass


threading.Thread(target=sampler, daemon=True).start()

# ---------------------------------------------------------------- helpers ----


def overview():
    vm = psutil.virtual_memory()
    sw = psutil.swap_memory()
    batt = None
    try:
        b = psutil.sensors_battery()
        if b:
            batt = {"percent": round(b.percent), "plugged": b.power_plugged,
                    "secsleft": b.secsleft if b.secsleft and b.secsleft > 0 else None}
    except Exception:
        pass
    temps = []
    try:
        for name, entries in (psutil.sensors_temperatures() or {}).items():
            for e in entries:
                if e.current:
                    temps.append({"label": e.label or name, "c": round(e.current)})
    except Exception:
        pass
    la = os.getloadavg()
    return {
        "hostname": os.uname().nodename,
        "os": " ".join(os.uname().sysname.split()) + " " + os.uname().release,
        "uptime": int(time.time() - BOOT),
        "cpu": psutil.cpu_percent(interval=None),
        "percore": psutil.cpu_percent(interval=None, percpu=True),
        "cores": psutil.cpu_count(logical=True),
        "freq": (psutil.cpu_freq().current if psutil.cpu_freq() else None),
        "load": [round(x, 2) for x in la],
        "mem": {"total": vm.total, "used": vm.total - vm.available,
                "available": vm.available, "percent": vm.percent},
        "swap": {"total": sw.total, "used": sw.used, "percent": sw.percent},
        "battery": batt,
        "temps": temps[:6],
        "nproc": len(psutil.pids()),
    }


def disks():
    out = []
    seen = set()
    for p in psutil.disk_partitions(all=False):
        if p.device in seen or p.fstype in ("squashfs", "tmpfs", "devtmpfs"):
            continue
        seen.add(p.device)
        try:
            u = psutil.disk_usage(p.mountpoint)
        except OSError:
            continue
        out.append({"device": p.device, "mount": p.mountpoint, "fstype": p.fstype,
                    "total": u.total, "used": u.used, "free": u.free,
                    "percent": u.percent})
    out.sort(key=lambda d: -d["total"])
    return out


def processes(sort="cpu", query=""):
    procs = []
    me = os.getuid()
    q = query.lower()
    for p in psutil.process_iter(["pid", "name", "username", "memory_info",
                                  "cpu_percent", "status", "create_time",
                                  "cmdline", "uids"]):
        try:
            info = p.info
            name = info["name"] or "?"
            cmd = " ".join(info["cmdline"] or [])[:160]
            if q and q not in name.lower() and q not in cmd.lower():
                continue
            procs.append({
                "pid": info["pid"], "name": name,
                "user": info["username"] or "?",
                "mem": info["memory_info"].rss if info["memory_info"] else 0,
                "cpu": info["cpu_percent"] or 0.0,
                "status": info["status"],
                "started": info["create_time"],
                "cmd": cmd,
                "mine": bool(info["uids"] and info["uids"].real == me),
            })
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    key = "mem" if sort == "mem" else "cpu"
    procs.sort(key=lambda x: -x[key])
    return procs[:60]


def kill_proc(pid, force=False):
    p = psutil.Process(pid)
    name = p.name()
    if force:
        p.kill()
    else:
        p.terminate()
    return name


def users():
    sessions = []
    for u in psutil.users():
        sessions.append({"name": u.name, "terminal": u.terminal or "-",
                         "host": u.host or "local", "started": u.started})
    agg = {}
    for p in psutil.process_iter(["username", "memory_info", "cpu_percent"]):
        try:
            un = p.info["username"] or "?"
            a = agg.setdefault(un, {"user": un, "procs": 0, "mem": 0, "cpu": 0.0})
            a["procs"] += 1
            a["mem"] += p.info["memory_info"].rss if p.info["memory_info"] else 0
            a["cpu"] += p.info["cpu_percent"] or 0.0
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    top = sorted(agg.values(), key=lambda a: -a["mem"])[:15]
    human = []
    for u in pwd.getpwall():
        if 1000 <= u.pw_uid < 60000:
            human.append({"user": u.pw_name, "uid": u.pw_uid, "home": u.pw_dir,
                          "shell": u.pw_shell})
    return {"sessions": sessions, "usage": top, "accounts": human}


def browse(path):
    path = os.path.realpath(path or HOME)
    if not os.path.isdir(path):
        raise ValueError("not a directory")
    entries = []
    with os.scandir(path) as it:
        for e in it:
            try:
                st = e.stat(follow_symlinks=False)
                entries.append({
                    "name": e.name,
                    "dir": e.is_dir(follow_symlinks=False),
                    "link": e.is_symlink(),
                    "size": st.st_size,
                    "mtime": st.st_mtime,
                    "hidden": e.name.startswith("."),
                    "pv": preview_kind(e.name),
                })
            except OSError:
                continue
    entries.sort(key=lambda x: (not x["dir"], x["name"].lower()))
    return {"path": path, "parent": os.path.dirname(path) if path != "/" else None,
            "entries": entries[:800], "truncated": len(entries) > 800}


def dir_size(path, deadline):
    """Recursive size with a wall-clock deadline; returns (bytes, complete)."""
    total = 0
    complete = True
    stack = [path]
    while stack:
        if time.time() > deadline:
            return total, False
        d = stack.pop()
        try:
            with os.scandir(d) as it:
                for e in it:
                    try:
                        if e.is_dir(follow_symlinks=False):
                            stack.append(e.path)
                        else:
                            total += e.stat(follow_symlinks=False).st_size
                    except OSError:
                        continue
        except OSError:
            complete = False
    return total, complete


def analyze(path):
    path = os.path.realpath(path or HOME)
    if not os.path.isdir(path):
        raise ValueError("not a directory")
    deadline = time.time() + 15
    items = []
    files_size = 0
    with os.scandir(path) as it:
        children = list(it)
    for e in children:
        try:
            if e.is_dir(follow_symlinks=False):
                size, complete = dir_size(e.path, deadline)
                items.append({"name": e.name, "dir": True, "size": size,
                              "complete": complete})
            else:
                files_size += e.stat(follow_symlinks=False).st_size
        except OSError:
            continue
    if files_size:
        items.append({"name": "(files here)", "dir": False, "size": files_size,
                      "complete": True})
    items.sort(key=lambda x: -x["size"])
    return {"path": path, "items": items[:40],
            "timedout": time.time() > deadline}


def open_file(path):
    path = os.path.realpath(path)
    if not os.path.exists(path):
        raise ValueError("no such file")
    subprocess.Popen(["xdg-open", path], stdout=subprocess.DEVNULL,
                     stderr=subprocess.DEVNULL,
                     env={**os.environ})
    return path


def trash_file(path):
    path = os.path.realpath(path)
    if not path.startswith(HOME + os.sep):
        raise ValueError("only files under your home can be trashed from here")
    if not os.path.exists(path):
        raise ValueError("no such file")
    if shutil.which("gio"):
        _run(["gio", "trash", path], check=True, capture_output=True)
    else:
        tdir = os.path.join(HOME, ".local/share/Trash/files")
        os.makedirs(tdir, exist_ok=True)
        shutil.move(path, os.path.join(tdir, os.path.basename(path) + "." +
                                       secrets.token_hex(4)))
    return path


CACHE_DIR = os.path.join(HOME, ".cache")


def quick_size(path, budget=8):
    if not os.path.exists(path):
        return 0
    size, _ = dir_size(path, time.time() + budget)
    return size


# ---- cleanup lenses: duplicates, and big files nobody has opened in years ---
# The filename index holds paths only, so both of these walk the chosen folder
# directly under a wall-clock deadline. Bounded work, honest partial results:
# `truncated` tells the UI the walk stopped early.

def _walk_files(root, deadline, min_size):
    root = os.path.realpath(os.path.expanduser(root or HOME))
    if not os.path.isdir(root):
        raise ValueError("not a folder: " + root)
    truncated = False
    for base, dirs, files in os.walk(root, topdown=True):
        if time.time() > deadline:
            truncated = True
            break
        dirs[:] = [d for d in dirs if d not in PRUNE_NAMES
                   and not os.path.islink(os.path.join(base, d))]
        for name in files:
            p = os.path.join(base, name)
            try:
                st = os.lstat(p)
            except OSError:
                continue
            if not stat.S_ISREG(st.st_mode) or st.st_size < min_size:
                continue
            yield p, st
    if truncated:
        yield None, None          # sentinel: the walk was cut short


def _digest(path, whole=False):
    h = hashlib.sha256()
    try:
        with open(path, "rb") as f:
            if whole:
                for chunk in iter(lambda: f.read(1 << 20), b""):
                    h.update(chunk)
            else:
                h.update(f.read(1 << 16))
    except OSError:
        return None
    return h.hexdigest()


def find_duplicates(path=None, min_mb=1, seconds=20):
    """Group identical files. Size first, then a 64 KB head digest, then a full
    digest — so most candidates are eliminated without being read in full."""
    deadline = time.time() + max(2, min(60, float(seconds)))
    min_size = max(1, int(float(min_mb) * 1024 * 1024))
    by_size, truncated = {}, False
    for p, st in _walk_files(path, deadline, min_size):
        if p is None:
            truncated = True
            break
        by_size.setdefault(st.st_size, []).append(p)
    groups = []
    for size, paths in sorted(by_size.items(), key=lambda kv: -kv[0]):
        if len(paths) < 2 or time.time() > deadline:
            continue
        by_head = {}
        for p in paths:
            d = _digest(p)
            if d:
                by_head.setdefault(d, []).append(p)
        for head, same_head in by_head.items():
            if len(same_head) < 2:
                continue
            by_full = {}
            for p in same_head:
                d = _digest(p, whole=True)
                if d:
                    by_full.setdefault(d, []).append(p)
            for full, dupes in by_full.items():
                if len(dupes) > 1:
                    groups.append({"size": size, "count": len(dupes),
                                   "wasted": size * (len(dupes) - 1),
                                   "paths": sorted(dupes)[:12]})
    groups.sort(key=lambda g: -g["wasted"])
    return {"groups": groups[:60], "truncated": truncated,
            "wasted": sum(g["wasted"] for g in groups),
            "root": os.path.realpath(os.path.expanduser(path or HOME))}


def find_old_large(path=None, min_mb=100, days=365, seconds=20):
    """Big files nothing has read in a long time — the other cleanup lens."""
    deadline = time.time() + max(2, min(60, float(seconds)))
    min_size = max(1, int(float(min_mb) * 1024 * 1024))
    cutoff = time.time() - max(1, float(days)) * 86400
    out, truncated = [], False
    for p, st in _walk_files(path, deadline, min_size):
        if p is None:
            truncated = True
            break
        last = max(st.st_atime, st.st_mtime)
        if last <= cutoff:
            out.append({"path": p, "size": st.st_size, "atime": st.st_atime,
                        "mtime": st.st_mtime})
    out.sort(key=lambda f: -f["size"])
    return {"files": out[:200], "truncated": truncated,
            "total": sum(f["size"] for f in out),
            "root": os.path.realpath(os.path.expanduser(path or HOME))}


def cleanup_report():
    deadline_each = 6
    targets = []
    thumbs = os.path.join(CACHE_DIR, "thumbnails")
    targets.append({"id": "thumbnails", "label": "Thumbnail cache",
                    "path": thumbs, "size": quick_size(thumbs, deadline_each),
                    "desc": "Image/video preview thumbnails. Regenerated on demand."})
    trash = os.path.join(HOME, ".local/share/Trash")
    targets.append({"id": "trash", "label": "Trash",
                    "path": trash, "size": quick_size(trash, deadline_each),
                    "desc": "Deleted files waiting in the trash bin."})
    pipc = _run(["python3", "-m", "pip", "cache", "dir"],
                          capture_output=True, text=True)
    pip_path = pipc.stdout.strip() if pipc.returncode == 0 else ""
    targets.append({"id": "pip", "label": "pip download cache",
                    "path": pip_path,
                    "size": quick_size(pip_path, deadline_each) if pip_path else 0,
                    "desc": "Cached Python package downloads."})
    # top offenders inside ~/.cache (deletable individually)
    cache_top = []
    try:
        with os.scandir(CACHE_DIR) as it:
            kids = [e for e in it if e.is_dir(follow_symlinks=False)]
        deadline = time.time() + 20
        for e in kids:
            s, _ = dir_size(e.path, deadline)
            cache_top.append({"name": e.name, "path": e.path, "size": s})
        cache_top.sort(key=lambda x: -x["size"])
        cache_top = cache_top[:12]
    except OSError:
        pass
    # sudo-only items: report size + the command to run yourself
    sudo_items = []
    j = _run(["journalctl", "--disk-usage"], capture_output=True,
                       text=True)
    if j.returncode == 0:
        sudo_items.append({"label": "systemd journal logs",
                           "info": j.stdout.strip(),
                           "cmd": "sudo journalctl --vacuum-size=200M"})
    apt = quick_size("/var/cache/apt/archives", 4)
    if apt > 50 * 2 ** 20:
        sudo_items.append({"label": "APT package cache",
                           "info": f"{apt / 2**30:.2f} GB in /var/cache/apt/archives",
                           "cmd": "sudo apt-get clean"})
    return {"targets": targets, "cache_top": cache_top, "sudo": sudo_items}


def do_clean(target, path=None):
    if target == "thumbnails":
        p = os.path.join(CACHE_DIR, "thumbnails")
        freed = quick_size(p)
        shutil.rmtree(p, ignore_errors=True)
        return freed
    if target == "trash":
        t = os.path.join(HOME, ".local/share/Trash")
        freed = quick_size(t)
        if shutil.which("gio"):
            _run(["gio", "trash", "--empty"], capture_output=True)
        for sub in ("files", "info"):
            d = os.path.join(t, sub)
            if os.path.isdir(d):
                for name in os.listdir(d):
                    fp = os.path.join(d, name)
                    if os.path.isdir(fp) and not os.path.islink(fp):
                        shutil.rmtree(fp, ignore_errors=True)
                    else:
                        try:
                            os.unlink(fp)
                        except OSError:
                            pass
        return freed
    if target == "pip":
        r = _run(["python3", "-m", "pip", "cache", "dir"],
                           capture_output=True, text=True)
        freed = quick_size(r.stdout.strip()) if r.returncode == 0 else 0
        _run(["python3", "-m", "pip", "cache", "purge"],
                       capture_output=True)
        return freed
    if target == "cachedir":
        p = os.path.realpath(path or "")
        if not p.startswith(CACHE_DIR + os.sep) or p == CACHE_DIR:
            raise ValueError("only directories inside ~/.cache can be removed here")
        freed = quick_size(p)
        shutil.rmtree(p, ignore_errors=True)
        return freed
    raise ValueError("unknown cleanup target")


# -------------------------------------------------------------- hardware -----


def _read1(path):
    try:
        with open(path) as f:
            return f.read().strip()
    except OSError:
        return None


def hw_info():
    dmi = "/sys/class/dmi/id/"
    model = " ".join(x for x in (_read1(dmi + "sys_vendor"),
                                 _read1(dmi + "product_family")) if x)
    cpu_model = None
    try:
        with open("/proc/cpuinfo") as f:
            for line in f:
                if line.startswith("model name"):
                    cpu_model = line.split(":", 1)[1].strip()
                    break
    except OSError:
        pass
    bat, b = {}, "/sys/class/power_supply/BAT0/"
    if os.path.isdir(b):
        full = _read1(b + "energy_full") or _read1(b + "charge_full")
        design = (_read1(b + "energy_full_design")
                  or _read1(b + "charge_full_design"))
        power = _read1(b + "power_now")
        bat = {
            "cycles": _read1(b + "cycle_count"),
            "health": (min(100, round(int(full) / int(design) * 100))
                       if full and design and int(design) else None),
            "status": _read1(b + "status"),
            "capacity": _read1(b + "capacity"),
            "power_w": round(int(power) / 1e6, 1) if power else None,
        }
    wifi = {}
    try:
        r = _run(["nmcli", "-t", "-f", "active,ssid,signal",
                            "dev", "wifi"], capture_output=True, text=True,
                           timeout=6)
        for line in r.stdout.splitlines():
            if line.startswith("yes"):
                parts = line.split(":")
                wifi = {"ssid": parts[1], "signal": parts[2] if len(parts) > 2
                        else None}
                break
    except Exception:  # noqa: BLE001
        pass
    try:
        with open("/proc/net/wireless") as f:
            last = f.read().splitlines()[-1].split()
            if len(last) > 3 and not last[0].startswith("Inter"):
                wifi["dbm"] = int(float(last[3]))
    except (OSError, ValueError, IndexError):
        pass
    bright = None
    for bl in glob.glob("/sys/class/backlight/*"):
        cur, mx = _read1(bl + "/brightness"), _read1(bl + "/max_brightness")
        if cur and mx and int(mx):
            bright = round(int(cur) / int(mx) * 100)
    return {"model": model or None,
            "product": _read1(dmi + "product_name"),
            "bios": _read1(dmi + "bios_version"),
            "cpu": cpu_model, "gpu": GPU_NAME,
            "ram_gb": round(psutil.virtual_memory().total / 2**30, 1),
            "battery": bat, "wifi": wifi, "brightness": bright}


# ---------------------------------------------------------------- search -----

INDEX_DIR = os.path.join(CACHE_DIR, "perch")
INDEX_FILE = os.path.join(INDEX_DIR, "index.txt")
INDEX_STATUS = {"state": "none", "count": 0, "built": 0, "error": None}
_index_lock = threading.Lock()
PRUNE_PATHS = {"/proc", "/sys", "/dev", "/run", "/snap", "/var/snap",
               "/var/lib/docker", "/lost+found", INDEX_DIR}
PRUNE_NAMES = {".git", "__pycache__"}


def build_index():
    with _index_lock:
        if INDEX_STATUS["state"] == "building":
            return
        INDEX_STATUS.update(state="building", error=None)
    try:
        os.makedirs(INDEX_DIR, exist_ok=True)
        tmp = INDEX_FILE + ".tmp"
        count = 0
        with open(tmp, "w", errors="replace") as f:
            for root, dirs, files in os.walk("/", topdown=True):
                dirs[:] = [d for d in dirs if d not in PRUNE_NAMES
                           and os.path.join(root, d) not in PRUNE_PATHS]
                for name in dirs:
                    f.write(os.path.join(root, name) + "\n")
                for name in files:
                    f.write(os.path.join(root, name) + "\n")
                count += len(dirs) + len(files)
        os.replace(tmp, INDEX_FILE)
        INDEX_STATUS.update(state="ready", count=count, built=time.time())
    except Exception as e:  # noqa: BLE001
        INDEX_STATUS.update(state="error", error=str(e))


def index_boot():
    if os.path.exists(INDEX_FILE):
        age = time.time() - os.path.getmtime(INDEX_FILE)
        n = int(_run(["wc", "-l", INDEX_FILE], capture_output=True,
                               text=True).stdout.split()[0])
        INDEX_STATUS.update(state="ready", count=n,
                            built=os.path.getmtime(INDEX_FILE))
        if age < 12 * 3600:
            return
    build_index()


threading.Thread(target=index_boot, daemon=True).start()


def search_files(q, limit=250, regex=False):
    q = q.strip()
    status = {"index": INDEX_STATUS["state"], "count": INDEX_STATUS["count"],
              "built": INDEX_STATUS["built"]}
    if len(q) < 2:
        return {"status": status, "results": []}
    if not os.path.exists(INDEX_FILE):
        return {"status": status, "results": [],
                "note": "index is still being built — try again in a minute"}
    if regex:
        try:
            pat = re.compile(q, re.I)
        except re.error as e:
            raise ValueError(f"invalid regex: {e}")
        grep_args = ["grep", "-i", "-E"]
    else:
        grep_args = ["grep", "-i", "-F"]
    r = _run([*grep_args, "-m", str(limit * 4), "--", q,
                        INDEX_FILE], capture_output=True, text=True, timeout=20)
    if regex and r.returncode == 2:
        raise ValueError("grep rejected that pattern (use POSIX extended "
                         "regex, e.g. \\.log$ or ^/etc/.*conf)")
    starts, contains, pathonly = [], [], []
    ql = q.lower()
    for p in r.stdout.splitlines():
        base = os.path.basename(p).lower()
        if regex:
            m = pat.search(base)
            if m and m.start() == 0:
                starts.append(p)
            elif m:
                contains.append(p)
            else:
                pathonly.append(p)
        elif base.startswith(ql):
            starts.append(p)
        elif ql in base:
            contains.append(p)
        else:
            pathonly.append(p)
    ranked = (starts + contains + pathonly)[:limit]
    results = []
    for p in ranked:
        try:
            st = os.stat(p, follow_symlinks=False)
            results.append({"path": p, "dir": os.path.isdir(p),
                            "size": st.st_size, "mtime": st.st_mtime})
        except OSError:
            continue  # stale index entry
    return {"status": status, "results": results,
            "truncated": len(r.stdout.splitlines()) >= limit * 4}


# --------------------------------------------------------------- network -----


def net_ports():
    listen, established = [], 0
    me = os.getuid()
    names = {}
    for c in psutil.net_connections(kind="inet"):
        if c.status == "ESTABLISHED":
            established += 1
        if c.status != "LISTEN":
            continue
        name, mine, user, cmd = None, False, None, None
        if c.pid:
            if c.pid not in names:
                try:
                    pr = psutil.Process(c.pid)
                    with pr.oneshot():
                        names[c.pid] = (pr.name(), pr.uids().real == me,
                                        pr.username(),
                                        " ".join(pr.cmdline())[:200])
                except psutil.Error:
                    names[c.pid] = ("?", False, None, None)
            name, mine, user, cmd = names[c.pid]
        listen.append({"port": c.laddr.port, "addr": c.laddr.ip,
                       "pid": c.pid, "name": name, "mine": mine,
                       "user": user, "cmd": cmd,
                       "self": c.pid == os.getpid(),
                       "public": c.laddr.ip in ("0.0.0.0", "::")})
    listen.sort(key=lambda x: x["port"])
    seen, uniq = set(), []
    for p in listen:
        k = (p["port"], p["addr"])
        if k not in seen:
            seen.add(k)
            uniq.append(p)
    ifaces = []
    stats = psutil.net_if_stats()
    for nic, addrs in psutil.net_if_addrs().items():
        if nic == "lo":
            continue
        ips = [a.address for a in addrs if a.family.name == "AF_INET"]
        if ips:
            up = stats.get(nic) and stats[nic].isup
            ifaces.append({"nic": nic, "ips": ips, "up": bool(up)})
    return {"listen": uniq, "established": established, "ifaces": ifaces}


# ---- firewall ----
# Reading the actual rule set needs root everywhere (ufw/nft/iptables all
# refuse otherwise), so unprivileged we report what *is* readable — whether the
# service runs and ufw's own config flag — and offer the rule dump as a pkexec
# job, the same pattern as package installs.

def firewall_status():
    tools = {}
    for name, unit in (("ufw", "ufw"), ("firewalld", "firewalld"),
                       ("nftables", "nftables")):
        if not (shutil.which(name) or shutil.which(name.rstrip("d"))):
            continue
        r = _run(["systemctl", "is-active", unit], capture_output=True,
                 text=True, timeout=8)
        tools[name] = {"service": r.stdout.strip() or "unknown"}
    enabled = None
    try:
        with open("/etc/ufw/ufw.conf") as f:
            for line in f:
                if line.strip().startswith("ENABLED="):
                    enabled = line.strip().split("=", 1)[1].strip().lower() == "yes"
    except OSError:
        pass
    if "ufw" in tools:
        tools["ufw"]["enabled"] = enabled
    return {"tools": tools, "ufw_enabled": enabled,
            "any": bool(tools),
            "can_dump": bool(shutil.which("ufw") or shutil.which("nft")
                             or shutil.which("iptables"))}


def firewall_rules_job():
    """Dump the live rule set — needs root, so it runs as a privileged job."""
    if shutil.which("ufw"):
        argv, title = ["ufw", "status", "verbose"], "Firewall rules (ufw)"
    elif shutil.which("nft"):
        argv, title = ["nft", "list", "ruleset"], "Firewall rules (nftables)"
    elif shutil.which("iptables"):
        argv, title = ["iptables", "-L", "-n", "-v"], "Firewall rules (iptables)"
    else:
        raise ValueError("no firewall tool found (looked for ufw, nft, iptables)")
    return start_job(argv, title, privileged=True)


# ---- disk health ----

def disk_health():
    """Physical devices from /sys/block (unprivileged). SMART itself needs
    root, so the detailed report is a separate privileged job."""
    out = []
    try:
        names = sorted(os.listdir("/sys/block"))
    except OSError:
        return {"devices": [], "smartctl": False}
    for name in names:
        if name.startswith(("loop", "ram", "dm-", "zram", "sr")):
            continue
        base = "/sys/block/" + name
        def rd(rel, default=""):
            try:
                with open(os.path.join(base, rel)) as f:
                    return f.read().strip()
            except OSError:
                return default
        sectors = rd("size", "0")
        try:
            size = int(sectors) * 512
        except ValueError:
            size = 0
        model = rd("device/model") or rd("device/name")
        out.append({
            "name": name, "size": size,
            "model": model, "vendor": rd("device/vendor"),
            "rotational": rd("queue/rotational") == "1",
            "scheduler": rd("queue/scheduler"),
            "readonly": rd("ro") == "1",
        })
    return {"devices": out, "smartctl": bool(shutil.which("smartctl"))}


SMART_DEV_RE = re.compile(r"[a-zA-Z0-9]{1,20}")


def smart_job(device):
    if not shutil.which("smartctl"):
        raise ValueError("smartctl not installed — install the smartmontools "
                         "package to read drive health")
    if not SMART_DEV_RE.fullmatch(device or ""):
        raise ValueError("bad device name")
    return start_job(["smartctl", "-H", "-A", "/dev/" + device],
                     f"SMART health for /dev/{device}", privileged=True)


def kill_port(port, force=False, pid=None):
    """Stop whatever is listening on `port`.

    `pid` (optional) pins the kill to the process the UI actually showed, so a
    listener that died and had its port re-bound in between is never hit by
    mistake.
    """
    for c in psutil.net_connections(kind="inet"):
        if c.status != "LISTEN" or c.laddr.port != port or not c.pid:
            continue
        if pid and c.pid != pid:
            continue
        name = kill_proc(c.pid, force)
        return f"{name} (pid {c.pid})"
    if pid:
        raise ValueError(f"pid {pid} is no longer listening on port {port} "
                         "— refresh the list")
    raise ValueError(f"no killable process found on port {port} "
                     "(root-owned listeners can't be ended from here)")


# ---------------------------------------------------------------- docker -----


def _docker_json(args, timeout=10):
    r = _run(["docker", *args], capture_output=True, text=True,
                       timeout=timeout)
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or "docker unavailable")[:300])
    return [json.loads(line) for line in r.stdout.splitlines() if line.strip()]


def docker_info():
    containers = [{"id": c["ID"], "name": c["Names"], "image": c["Image"],
                   "state": c["State"], "status": c["Status"],
                   "ports": c.get("Ports", "")}
                  for c in _docker_json(["ps", "-a", "--format", "{{json .}}"])]
    images = [{"repo": i["Repository"], "tag": i["Tag"], "size": i["Size"],
               "id": i["ID"]}
              for i in _docker_json(["images", "--format", "{{json .}}"])][:20]
    return {"containers": containers, "images": images}


DOCKER_ACTIONS = {"start", "stop", "restart", "rm"}


def docker_action(cid, action):
    if action not in DOCKER_ACTIONS or not re.fullmatch(r"[0-9a-f]{4,64}", cid):
        raise ValueError("bad docker action")
    r = _run(["docker", action, cid], capture_output=True, text=True,
                       timeout=60)
    if r.returncode != 0:
        raise ValueError(r.stderr.strip()[:300])
    return r.stdout.strip()


def docker_logs(cid):
    if not re.fullmatch(r"[0-9a-f]{4,64}", cid):
        raise ValueError("bad container id")
    r = _run(["docker", "logs", "--tail", "150", cid],
                       capture_output=True, text=True, timeout=15)
    return {"logs": (r.stdout + r.stderr)[-20000:]}


# ------------------------------------------------- other container engines ---
# Docker is handled above; this section detects whatever *else* is installed
# (Podman, nerdctl, LXD, Kubernetes) and lists its containers/pods with the
# same shape, so the Dev tab can render them all through one code path.

CTR_ENGINES = ("docker", "podman", "nerdctl")
CTR_ACTIONS = {"start", "stop", "restart", "rm"}
CTR_ID_RE = re.compile(r"[0-9a-zA-Z][\w.-]{0,63}")
K8S_NAME_RE = re.compile(r"[a-z0-9][a-z0-9.-]{0,252}")


def _ctr_ports(ports):
    """Normalise a ports field: docker gives a string, podman a list."""
    if not isinstance(ports, list):
        return str(ports or "")
    out = []
    for p in ports:
        if not isinstance(p, dict):
            out.append(str(p))
            continue
        host = p.get("host_port") or p.get("hostPort") or ""
        cont = p.get("container_port") or p.get("containerPort") or ""
        proto = p.get("protocol") or "tcp"
        out.append(f"{host}->{cont}/{proto}" if host else f"{cont}/{proto}")
    return ", ".join(str(x) for x in out)


def _ctr_norm(c):
    names = c.get("Names") or c.get("Name") or ""
    if isinstance(names, list):
        names = ", ".join(names)
    state = str(c.get("State") or "").lower()
    status = c.get("Status") or ""
    if state in ("", "unknown"):
        state = "running" if str(status).lower().startswith("up") else state
    return {"id": str(c.get("ID") or c.get("Id") or "")[:64],
            "name": names, "image": c.get("Image") or "",
            "state": state, "status": str(status),
            "ports": _ctr_ports(c.get("Ports"))}


def _ctr_ps(engine):
    """`<engine> ps -a` as a normalised list.

    Docker/nerdctl honour the Go template and emit one JSON object per line;
    Podman ignores it and emits a single JSON array — both are accepted.
    """
    r = _run([engine, "ps", "-a", "--format", "{{json .}}"],
             capture_output=True, text=True, timeout=12)
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or engine + " unavailable")[:300])
    text = (r.stdout or "").strip()
    if not text:
        return []
    rows = (json.loads(text) if text[0] == "["
            else [json.loads(x) for x in text.splitlines() if x.strip()])
    return [_ctr_norm(c) for c in rows][:200]


def _ctr_version(engine):
    r = _run([engine, "--version"], capture_output=True, text=True, timeout=8)
    lines = ((r.stdout or "") + (r.stderr or "")).strip().splitlines()
    return lines[0][:60] if lines else ""


def _lxd_containers():
    """LXD/Incus instances, if that client is installed and talking to a daemon."""
    for binary in ("incus", "lxc"):
        if not shutil.which(binary):
            continue
        r = _run([binary, "list", "--format", "json"],
                 capture_output=True, text=True, timeout=12)
        if r.returncode != 0 or not (r.stdout or "").strip():
            continue
        try:
            rows = json.loads(r.stdout)
        except ValueError:
            continue
        out = []
        for i in rows[:200]:
            # a stopped instance reports "state": null, so `or {}` at every hop
            net = (i.get("state") or {}).get("network") or {}
            ips = [a["address"] for iface, d in net.items() if iface != "lo"
                   for a in (d.get("addresses") or [])
                   if a.get("family") == "inet"]
            out.append({"id": i.get("name", ""), "name": i.get("name", ""),
                        "image": ((i.get("config") or {})
                                  .get("image.description") or "")[:60],
                        "state": str(i.get("status", "")).lower(),
                        "status": i.get("status", ""), "ports": ", ".join(ips)})
        return {"engine": binary, "kind": "lxd", "version": _ctr_version(binary),
                "containers": out, "error": None}
    return None


def _k8s_pods():
    """Pods from the current kubectl context, or None when there isn't one."""
    if not shutil.which("kubectl"):
        return None
    r = _run(["kubectl", "config", "current-context"],
             capture_output=True, text=True, timeout=6)
    if r.returncode != 0:
        return None
    ctx = r.stdout.strip()
    # --request-timeout keeps an unreachable cluster from stalling the panel
    r = _run(["kubectl", "get", "pods", "--all-namespaces", "-o", "json",
              "--request-timeout=8s"], capture_output=True, text=True,
             timeout=20)
    if r.returncode != 0:
        return {"context": ctx, "pods": [],
                "error": (r.stderr.strip() or "cluster unreachable")[:300]}
    try:
        items = json.loads(r.stdout or "{}").get("items", [])
    except ValueError:
        return {"context": ctx, "pods": [], "error": "unreadable kubectl output"}
    pods = []
    for it in items[:300]:
        md, st = it.get("metadata") or {}, it.get("status") or {}
        cs = st.get("containerStatuses") or []
        pods.append({"ns": md.get("namespace", ""), "name": md.get("name", ""),
                     "phase": st.get("phase", ""),
                     "ready": f"{sum(1 for c in cs if c.get('ready'))}/{len(cs)}",
                     "restarts": sum(c.get("restartCount") or 0 for c in cs),
                     "node": (it.get("spec") or {}).get("nodeName", ""),
                     "ip": st.get("podIP") or ""})
    pods.sort(key=lambda p: (p["ns"], p["name"]))
    return {"context": ctx, "pods": pods, "error": None}


def container_envs():
    """Every container environment present on this machine, Docker included."""
    envs = []
    for engine in CTR_ENGINES:
        if not shutil.which(engine):
            continue
        env = {"engine": engine, "kind": "container",
               "version": _ctr_version(engine), "containers": [], "error": None}
        try:
            env["containers"] = _ctr_ps(engine)
        except Exception as e:  # noqa: BLE001 — one dead engine must not hide the rest
            env["error"] = str(e)[:300]
        envs.append(env)
    # one flaky/odd-shaped environment must not take the whole panel down
    try:
        lxd = _lxd_containers()
        if lxd:
            envs.append(lxd)
    except Exception:  # noqa: BLE001
        pass
    try:
        k8s = _k8s_pods()
    except Exception:  # noqa: BLE001
        k8s = None
    return {"envs": envs, "k8s": k8s}


def ctr_action(engine, cid, action):
    if engine not in CTR_ENGINES or action not in CTR_ACTIONS:
        raise ValueError("bad container action")
    if not CTR_ID_RE.fullmatch(cid or ""):
        raise ValueError("bad container id")
    r = _run([engine, action, cid], capture_output=True, text=True, timeout=60)
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or "command failed")[:300])
    return r.stdout.strip()


def ctr_logs(engine, cid, ns=""):
    """Recent log lines from a container (docker/podman/nerdctl) or a k8s pod."""
    if engine == "k8s":
        if not (K8S_NAME_RE.fullmatch(cid or "")
                and K8S_NAME_RE.fullmatch(ns or "")):
            raise ValueError("bad pod name")
        cmd = ["kubectl", "logs", "-n", ns, cid, "--tail", "150",
               "--all-containers=true"]
    else:
        if engine not in CTR_ENGINES:
            raise ValueError("unknown container engine")
        if not CTR_ID_RE.fullmatch(cid or ""):
            raise ValueError("bad container id")
        cmd = [engine, "logs", "--tail", "150", cid]
    r = _run(cmd, capture_output=True, text=True, timeout=20)
    return {"logs": (r.stdout + r.stderr)[-20000:]}


# -------------------------------------------------------------- services -----

SVC_RE = re.compile(r"^[\w@.\\:-]+\.service$")


def services():
    def list_units(scope, extra=()):
        r = _run(["systemctl", scope, "list-units", "--type=service",
                            "--all", "--no-pager", "--output=json", *extra],
                           capture_output=True, text=True, timeout=10)
        if r.returncode != 0:
            return []
        return json.loads(r.stdout or "[]")

    # whether each unit starts at login, so the UI can offer enable/disable
    enabled = {}
    r = _run(["systemctl", "--user", "list-unit-files", "--type=service",
              "--no-pager", "--output=json"],
             capture_output=True, text=True, timeout=10)
    if r.returncode == 0:
        try:
            for f in json.loads(r.stdout or "[]"):
                enabled[f.get("unit_file", "")] = f.get("state", "")
        except ValueError:
            pass
    user = [{"name": u["unit"], "active": u["active"], "sub": u["sub"],
             "desc": u["description"],
             "enabled": enabled.get(u["unit"], "")} for u in list_units("--user")]
    user.sort(key=lambda u: (u["active"] != "active", u["name"]))
    failed = [{"name": u["unit"], "desc": u["description"]}
              for u in list_units("--system", ("--state=failed",))]
    return {"user": user, "failed_system": failed}


SVC_ACTIONS = ("start", "stop", "restart", "enable", "disable")


def service_action(name, action):
    """Act on a user unit. enable/disable change whether it starts at login."""
    if action not in SVC_ACTIONS or not SVC_RE.fullmatch(name):
        raise ValueError("bad service action")
    r = _run(["systemctl", "--user", action, name],
                       capture_output=True, text=True, timeout=30)
    if r.returncode != 0:
        raise ValueError(r.stderr.strip()[:300])
    return "ok"


# -------------------------------------------------------------- dev info -----

TOOLS = [
    ("git", ["git", "--version"]), ("python", ["python3", "--version"]),
    ("pip", ["python3", "-m", "pip", "--version"]),
    ("node", ["node", "--version"]), ("npm", ["npm", "--version"]),
    ("docker", ["docker", "--version"]),
    ("compose", ["docker", "compose", "version"]),
    ("go", ["go", "version"]), ("rustc", ["rustc", "--version"]),
    ("java", ["java", "-version"]), ("gcc", ["gcc", "--version"]),
    ("make", ["make", "--version"]), ("code", ["code", "--version"]),
]


def devinfo():
    tools = []
    for label, cmd in TOOLS:
        if not shutil.which(cmd[0]):
            continue
        try:
            r = _run(cmd, capture_output=True, text=True, timeout=8)
            line = (r.stdout or r.stderr).strip().splitlines()
            if line:
                tools.append({"tool": label, "version": line[0][:80]})
        except Exception:  # noqa: BLE001
            continue
    git_id = {}
    for k in ("user.name", "user.email"):
        r = _run(["git", "config", "--global", k],
                           capture_output=True, text=True)
        git_id[k] = r.stdout.strip()
    return {"tools": tools, "git": git_id}


def open_editor(path):
    path = os.path.realpath(path)
    for ed in ("code", "subl", "gedit"):
        if shutil.which(ed):
            subprocess.Popen([ed, path], stdout=subprocess.DEVNULL,
                             stderr=subprocess.DEVNULL)
            return ed
    raise ValueError("no editor found (looked for code, subl, gedit)")


# ---------------------------------------------------------------- kernel -----

TUNABLES = [
    ("vm.swappiness", "How eagerly the kernel swaps RAM to disk (0–200). "
     "Lower = keep more in RAM; 10 is a common desktop choice."),
    ("vm.vfs_cache_pressure", "Willingness to drop directory/inode caches. "
     "Lower keeps filesystem metadata cached (snappier file ops)."),
    ("vm.dirty_ratio", "Max % of RAM holding unwritten (dirty) data before "
     "writers are blocked to flush."),
    ("vm.dirty_background_ratio", "% of RAM with dirty data at which background "
     "flushing to disk starts."),
    ("vm.overcommit_memory", "0 heuristic / 1 always allow / 2 strict. "
     "Memory-allocation overcommit policy."),
    ("vm.max_map_count", "Max memory-map areas per process. Elasticsearch/"
     "some games need this raised (e.g. 262144)."),
    ("fs.file-max", "System-wide limit on open file handles."),
    ("fs.inotify.max_user_watches", "Max files watched for changes per user. "
     "VS Code / webpack watchers hit this — raise to 524288 if you see "
     "ENOSPC watch errors."),
    ("fs.inotify.max_user_instances", "Max inotify instances per user."),
    ("net.core.somaxconn", "Max pending-connection backlog per listening "
     "socket. Raise for busy local API servers."),
    ("net.ipv4.ip_local_port_range", "Ephemeral (outgoing) port range."),
    ("kernel.pid_max", "Highest PID before wrap-around."),
]


def _sysctl_read(key):
    try:
        with open("/proc/sys/" + key.replace(".", "/")) as f:
            return " ".join(f.read().split())
    except OSError:
        return None


def kernel_info():
    tunables = []
    for key, desc in TUNABLES:
        val = _sysctl_read(key)
        if val is not None:
            tunables.append({"key": key, "value": val, "desc": desc})
    governors = []
    import glob
    for pol in sorted(glob.glob("/sys/devices/system/cpu/cpufreq/policy*")):
        try:
            with open(os.path.join(pol, "scaling_governor")) as f:
                gov = f.read().strip()
            with open(os.path.join(pol, "scaling_available_governors")) as f:
                avail = f.read().split()
            governors.append({"policy": os.path.basename(pol), "governor": gov,
                              "available": avail})
        except OSError:
            continue
    thp = None
    try:
        with open("/sys/kernel/mm/transparent_hugepage/enabled") as f:
            thp = f.read().strip()
    except OSError:
        pass
    with open("/proc/version") as f:
        version = f.read().strip()
    with open("/proc/cmdline") as f:
        cmdline = f.read().strip()
    mods = []
    try:
        with open("/proc/modules") as f:
            for line in f:
                p = line.split()
                mods.append({"name": p[0], "size": int(p[1]),
                             "used_by": p[3] if p[3] != "-" else ""})
        mods.sort(key=lambda m: -m["size"])
        mods = mods[:14]
    except OSError:
        pass
    return {"version": version, "cmdline": cmdline, "tunables": tunables,
            "governors": governors, "thp": thp, "modules": mods,
            "nmodules": sum(1 for _ in open("/proc/modules"))}


# ------------------------------------------------------------------- logs ----

LOG_SOURCES = {"system": ["--system"], "user": ["--user"],
               "kernel": ["--system", "-k"]}


def read_logs(source="system", n=150, q="", prio="", cursor=""):
    args = ["journalctl", *LOG_SOURCES.get(source, ["--system"]),
            "-o", "json", "--no-pager"]
    if cursor:
        args += ["--after-cursor", cursor]
    else:
        args += ["-n", str(min(int(n or 150), 500))]
    if q:
        args += ["-g", q]
    if prio:
        args += ["-p", prio]
    r = _run(args, capture_output=True, text=True, timeout=20)
    if r.returncode != 0 and not r.stdout:
        raise ValueError((r.stderr.strip() or "journalctl failed")[:200])
    entries, last_cursor = [], cursor
    for line in r.stdout.splitlines()[-500:]:
        try:
            e = json.loads(line)
        except json.JSONDecodeError:
            continue
        msg = e.get("MESSAGE", "")
        if isinstance(msg, list):  # journald stores non-UTF8 as byte arrays
            msg = bytes(msg).decode("utf-8", "replace")
        ts = int(e.get("__REALTIME_TIMESTAMP", 0)) / 1e6
        entries.append({
            "t": ts,
            "unit": (e.get("_SYSTEMD_UNIT") or e.get("SYSLOG_IDENTIFIER")
                     or e.get("_COMM") or "?"),
            "msg": str(msg)[:600],
            "prio": int(e.get("PRIORITY", 6)),
        })
        last_cursor = e.get("__CURSOR", last_cursor)
    return {"entries": entries, "cursor": last_cursor}


# --------------------------------------------------------------------- ai ----

AI_LOCK = threading.Lock()
AI_SESSION = {"id": None}
CLAUDE_BIN = shutil.which("claude") or os.path.expanduser("~/.local/bin/claude")


def ai_snapshot():
    o = overview()
    lines = [
        f"Host: {o['hostname']}, {o['os']}, uptime {o['uptime']//3600}h",
        f"CPU {o['cpu']:.0f}% (load {o['load']}), "
        f"mem {o['mem']['percent']:.0f}% used "
        f"({o['mem']['used']/2**30:.1f}/{o['mem']['total']/2**30:.1f} GB), "
        f"swap {o['swap']['used']/2**30:.1f} GB",
    ]
    if GPU_CARD and HISTORY:
        lines.append(f"GPU busy {HISTORY[-1].get('gpu') or 0:.0f}%, "
                     f"{_gpu_freq('cur')} MHz")
    for d in disks()[:3]:
        lines.append(f"Disk {d['mount']}: {d['percent']:.0f}% used, "
                     f"{d['free']/2**30:.0f} GB free")
    top = processes("cpu")[:6]
    lines.append("Top CPU: " + ", ".join(
        f"{p['name']}({p['cpu']:.0f}%/{p['mem']/2**20:.0f}MB)" for p in top))
    top = processes("mem")[:6]
    lines.append("Top RAM: " + ", ".join(
        f"{p['name']}({p['mem']/2**20:.0f}MB)" for p in top))
    ports = [f"{p['port']}({p['name'] or '?'})" for p in net_ports()["listen"]]
    lines.append("Listening ports: " + ", ".join(ports[:15]))
    return "\n".join(lines)


def ask_ai(prompt, include_snapshot=True, reset=False, extra=""):
    if not os.path.exists(CLAUDE_BIN):
        raise ValueError("claude CLI not found")
    if not AI_LOCK.acquire(blocking=False):
        raise ValueError("an AI request is already running — wait for it")
    try:
        if reset:
            AI_SESSION["id"] = None
        full = prompt.strip()
        if extra:
            full += "\n\n--- attached data ---\n" + extra[:8000]
        if include_snapshot and not AI_SESSION["id"]:
            full = ("You are the assistant inside a local Linux system "
                    "dashboard. Current system snapshot:\n" + ai_snapshot() +
                    "\n\nUser question: " + full)
        args = [CLAUDE_BIN, "-p", "--output-format", "json"]
        if AI_SESSION["id"]:
            args += ["--resume", AI_SESSION["id"]]
        args.append(full)
        r = _run(args, capture_output=True, text=True, timeout=240,
                           cwd=HOME)
        if r.returncode != 0:
            raise ValueError((r.stderr.strip() or "claude failed")[:400])
        out = json.loads(r.stdout)
        AI_SESSION["id"] = out.get("session_id") or AI_SESSION["id"]
        return {"text": out.get("result", "").strip(),
                "session": AI_SESSION["id"],
                "cost": out.get("total_cost_usd")}
    finally:
        AI_LOCK.release()


def health_report():
    """Rich system context → prioritized AI recommendations (markdown)."""
    ctx = [ai_snapshot()]
    try:
        u = pkg_updates()
        ctx.append(f"Updates: {u['count']} pending ({u['security']} security)")
    except Exception:  # noqa: BLE001
        pass
    try:
        fs = services().get("failed_system", [])
        if fs:
            ctx.append("Failed services: "
                       + ", ".join(s["name"] for s in fs[:10]))
    except Exception:  # noqa: BLE001
        pass
    try:
        pub = [f"{p['port']}/{p['name'] or '?'}"
               for p in net_ports()["listen"] if p.get("public")]
        if pub:
            ctx.append("Network-visible ports: " + ", ".join(pub[:15]))
    except Exception:  # noqa: BLE001
        pass
    prompt = (
        "You are a Linux sysadmin reviewing this machine. Using ONLY the data "
        "below, give a concise health report as markdown: a one-line overall "
        "verdict, then a prioritized list (critical → nice-to-have) of concrete "
        "issues with the exact command to fix each. Skip anything that looks "
        "healthy. Be specific and brief.\n\n" + "\n".join(ctx))
    return {"text": llm_oneshot("", prompt)}


def health_score():
    """Deterministic plain-language scorecard with one-click fixes.

    Severities: crit (-20), warn (-10), info (-3). Actions the UI knows:
    upgrade (job), clean / updates / logs / dev (open that tab).
    """
    finds = []

    def add(sev, title, detail, action=None):
        finds.append({"sev": sev, "title": title, "detail": detail,
                      "action": action})

    for d in disks():
        if d["percent"] >= 92:
            add("crit", f"Disk almost full: {d['mount']}",
                f"{d['percent']:.0f}% used, "
                f"{d['free'] // 2**30} GB left. Things may stop working.",
                "clean")
        elif d["percent"] >= 85:
            add("warn", f"Disk getting full: {d['mount']}",
                f"{d['percent']:.0f}% used, {d['free'] // 2**30} GB left.",
                "clean")
    try:
        u = pkg_updates()
        if u["security"]:
            add("warn", f"{u['security']} security updates waiting",
                "Security fixes should be installed promptly.", "upgrade")
        elif u["count"]:
            add("info", f"{u['count']} software updates available",
                "Not urgent, but staying current avoids surprises.",
                "updates")
    except Exception:  # noqa: BLE001
        pass
    try:
        fs = services().get("failed_system", [])
        if fs:
            add("warn", f"{len(fs)} system service(s) failed",
                ", ".join(s["name"] for s in fs[:5]), "dev")
    except Exception:  # noqa: BLE001
        pass
    if os.path.exists("/var/run/reboot-required"):
        add("info", "Restart pending",
            "An update needs a reboot to fully apply.")
    vm = psutil.virtual_memory()
    sw = psutil.swap_memory()
    if vm.percent >= 90 and sw.percent >= 50:
        add("warn", "Memory pressure is high",
            f"RAM {vm.percent:.0f}% and swap {sw.percent:.0f}% used — the "
            "machine may feel slow. Close heavy apps or check Processes.")
    try:
        temps = psutil.sensors_temperatures()
        top = max((t.current for ts in temps.values() for t in ts),
                  default=0)
        if top >= 85:
            add("warn", f"Running hot: {top:.0f}°C",
                "Sustained high temperature. Check vents and heavy "
                "processes.")
    except (OSError, AttributeError):
        pass
    try:
        b = psutil.sensors_battery()
        full, design = 0, 0
        for supply in glob.glob("/sys/class/power_supply/BAT*"):
            try:
                with open(supply + "/energy_full") as f:
                    full = int(f.read())
                with open(supply + "/energy_full_design") as f:
                    design = int(f.read())
            except OSError:
                pass
        if b and design and full / design < 0.7:
            add("info", "Battery has aged",
                f"Holds {full * 100 // design}% of its original capacity. "
                "Expect shorter battery life.")
    except (OSError, RuntimeError):
        pass
    penalty = {"crit": 20, "warn": 10, "info": 3}
    score = max(0, 100 - sum(penalty[f["sev"]] for f in finds))
    verdict = ("Excellent" if score >= 90 else
               "Good" if score >= 75 else
               "Needs attention" if score >= 50 else "Poor")
    return {"score": score, "verdict": verdict, "findings": finds}


# --------------------------------------------------------- llm providers -----
# Perch ships no HTTP SDK, so provider calls use urllib against each provider's
# documented wire format. Anthropic: POST /v1/messages, x-api-key +
# anthropic-version: 2023-06-01, response.content is a list of blocks (text is
# on type=="text" blocks). OpenAI-compatible: POST {base}/chat/completions,
# Bearer auth, choices[0].message.content. Ollama: POST {base}/api/chat.

LLM_DEFAULTS = {"provider": "claude-cli", "model": "", "base_url": "",
                "api_key": ""}


def _llm_file():
    return os.path.join(CFG_DIR, "llm.json")


def llm_cfg():
    cfg = dict(LLM_DEFAULTS)
    try:
        with open(_llm_file()) as f:
            cfg.update(json.load(f))
    except (OSError, ValueError):
        pass
    return cfg


def llm_public():
    c = llm_cfg()
    return {"provider": c["provider"], "model": c["model"],
            "base_url": c["base_url"], "has_key": bool(c["api_key"]),
            "cli_available": os.path.exists(CLAUDE_BIN),
            "providers": ["claude-cli", "anthropic", "openai", "ollama"]}


def llm_save(body):
    os.makedirs(CFG_DIR, exist_ok=True)
    cfg = llm_cfg()
    for k in ("provider", "model", "base_url"):
        if k in body:
            cfg[k] = str(body[k]).strip()
    # keep existing key unless a new non-empty one is supplied
    if body.get("api_key"):
        cfg["api_key"] = str(body["api_key"]).strip()
    if body.get("clear_key"):
        cfg["api_key"] = ""
    atomic_write(_llm_file(), json.dumps(cfg))
    os.chmod(_llm_file(), 0o600)
    return llm_public()


def _http_post_json(url, headers, payload, timeout=240):
    import urllib.request as ur
    data = json.dumps(payload).encode()
    req = ur.Request(url, data=data, method="POST",
                     headers={"Content-Type": "application/json", **headers})
    try:
        with ur.urlopen(req, timeout=timeout) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "replace")[:300]
        raise ValueError(f"provider HTTP {e.code}: {detail}")
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"provider request failed: {e}")


def _provider_chat(cfg, system, messages, max_tokens=1024):
    """messages: [{'role','content'}]. Returns the assistant text."""
    p = cfg["provider"]
    if p == "anthropic":
        if not cfg["api_key"]:
            raise ValueError("no Anthropic API key configured")
        base = cfg["base_url"] or "https://api.anthropic.com"
        payload = {"model": cfg["model"] or "claude-sonnet-5",
                   "max_tokens": max_tokens, "messages": messages}
        if system:
            payload["system"] = system
        out = _http_post_json(base.rstrip("/") + "/v1/messages",
                              {"x-api-key": cfg["api_key"],
                               "anthropic-version": "2023-06-01"}, payload)
        return "".join(b.get("text", "") for b in out.get("content", [])
                       if b.get("type") == "text").strip()
    if p == "openai":
        base = cfg["base_url"] or "https://api.openai.com/v1"
        msgs = ([{"role": "system", "content": system}] if system else []) \
            + messages
        payload = {"model": cfg["model"] or "gpt-4o-mini", "messages": msgs,
                   "max_tokens": max_tokens}
        hdr = {"Authorization": f"Bearer {cfg['api_key']}"} \
            if cfg["api_key"] else {}
        out = _http_post_json(base.rstrip("/") + "/chat/completions", hdr,
                              payload)
        return out["choices"][0]["message"]["content"].strip()
    if p == "ollama":
        base = cfg["base_url"] or "http://localhost:11434"
        msgs = ([{"role": "system", "content": system}] if system else []) \
            + messages
        out = _http_post_json(base.rstrip("/") + "/api/chat",
                              {}, {"model": cfg["model"] or "llama3.2",
                                   "messages": msgs, "stream": False})
        return out.get("message", {}).get("content", "").strip()
    raise ValueError(f"provider '{p}' has no HTTP path")


def llm_oneshot(system, user_text, max_tokens=1500):
    """Single-turn completion via the configured provider (used by health)."""
    cfg = llm_cfg()
    if cfg["provider"] == "claude-cli":
        return ask_ai((system + "\n\n" + user_text) if system else user_text,
                      include_snapshot=False, reset=True)["text"]
    return _provider_chat(cfg, system,
                          [{"role": "user", "content": user_text}], max_tokens)


def llm_test():
    cfg = llm_cfg()
    if cfg["provider"] == "claude-cli":
        if not os.path.exists(CLAUDE_BIN):
            raise ValueError("claude CLI not found")
        return {"ok": True, "reply": "claude CLI present"}
    reply = _provider_chat(cfg, "You are a connectivity test.",
                           [{"role": "user",
                             "content": "Reply with exactly: OK"}], 20)
    return {"ok": True, "reply": reply[:200]}


# ------------------------------------------------------ http collections -----

def _http_file():
    return os.path.join(CFG_DIR, "http.json")


def _http_store():
    try:
        with open(_http_file()) as f:
            s = json.load(f)
    except (OSError, ValueError):
        s = {}
    s.setdefault("collections", [])
    s.setdefault("environments", {})
    s.setdefault("active_env", "")
    s.setdefault("history", [])
    s.setdefault("flows", [])
    return s


def http_store_get():
    return _http_store()


def http_store_save(body):
    os.makedirs(CFG_DIR, exist_ok=True)
    s = _http_store()
    for k in ("collections", "environments", "active_env", "flows"):
        if k in body:
            s[k] = body[k]
    atomic_write(_http_file(), json.dumps(s))
    return s


def http_history_add(entry):
    s = _http_store()
    s["history"].insert(0, entry)
    s["history"] = s["history"][:50]
    os.makedirs(CFG_DIR, exist_ok=True)
    atomic_write(_http_file(), json.dumps(s))


# ------------------------------------------------------- git projects --------

GIT_SKIP = {"node_modules", ".cache", "snap", ".local", ".cargo", ".rustup",
            ".npm", ".venv", "venv", "__pycache__", ".git"}


def _git(path, *args, timeout=15):
    return _run(["git", "-C", path, *args], capture_output=True,
                          text=True, timeout=timeout)


def git_repos():
    repos, deadline = [], time.time() + 12
    roots = [HOME]
    for base in roots:
        for root, dirs, _ in os.walk(base):
            if time.time() > deadline:
                break
            depth = root[len(base):].count(os.sep)
            if depth > 3:
                dirs[:] = []
                continue
            dirs[:] = [d for d in dirs if d not in GIT_SKIP
                       and not (depth == 0 and d.startswith("."))]
            if ".git" in os.listdir(root) if os.path.isdir(root) else False:
                pass
            if os.path.isdir(os.path.join(root, ".git")):
                repos.append(root)
                dirs[:] = []  # don't descend into a repo
    out = []
    for path in sorted(repos)[:60]:
        try:
            branch = _git(path, "rev-parse", "--abbrev-ref", "HEAD").stdout.strip()
            porcelain = _git(path, "status", "--porcelain").stdout.splitlines()
            ahead = behind = 0
            rb = _git(path, "rev-list", "--left-right", "--count",
                      "@{upstream}...HEAD")
            if rb.returncode == 0 and rb.stdout.strip():
                behind, ahead = (int(x) for x in rb.stdout.split())
            last = _git(path, "log", "-1", "--format=%cr|%s").stdout.strip()
            out.append({
                "path": path,
                "name": os.path.relpath(path, HOME),
                "branch": branch or "?",
                "dirty": len(porcelain),
                "ahead": ahead, "behind": behind,
                "last": last,
            })
        except Exception:  # noqa: BLE001
            continue
    out.sort(key=lambda r: (-r["dirty"], -r["ahead"], r["name"]))
    return {"repos": out}


def git_action(path, action):
    path = os.path.realpath(path)
    if not os.path.isdir(os.path.join(path, ".git")) \
            or not path.startswith(HOME):
        raise ValueError("not a git repo in your home")
    if action == "fetch":
        return start_job(["git", "-C", path, "fetch", "--all", "--prune"],
                         f"git fetch — {os.path.basename(path)}")
    if action == "pull":
        return start_job(["git", "-C", path, "pull", "--ff-only"],
                         f"git pull — {os.path.basename(path)}")
    if action == "stash":
        return start_job(["git", "-C", path, "stash", "push", "-u"],
                         f"git stash — {os.path.basename(path)}")
    raise ValueError("unknown git action")


def project_scripts(path):
    """Runnable tasks a repo declares — npm/yarn/pnpm scripts, Make targets,
    common Python entrypoints — so the Git tab can launch them."""
    path = os.path.realpath(path)
    if not path.startswith(HOME) or not os.path.isdir(path):
        raise ValueError("path not in your home")
    scripts = []
    pj = os.path.join(path, "package.json")
    if os.path.isfile(pj):
        try:
            with open(pj) as f:
                data = json.load(f)
            runner = ("pnpm" if os.path.exists(os.path.join(path, "pnpm-lock.yaml"))
                      else "yarn" if os.path.exists(os.path.join(path, "yarn.lock"))
                      else "npm")
            for name in list(data.get("scripts", {}))[:30]:
                scripts.append({"kind": runner, "name": name,
                                "cmd": f"{runner} run {name}"})
        except (OSError, ValueError):
            pass
    mk = next((m for m in ("Makefile", "makefile", "GNUmakefile")
               if os.path.isfile(os.path.join(path, m))), None)
    if mk:
        with open(os.path.join(path, mk)) as f:
            for line in f:
                m = re.match(r"^([a-zA-Z][\w.-]*)\s*:(?!=)", line)
                if m and m.group(1) not in ("PHONY",):
                    tgt = m.group(1)
                    if tgt not in [s["name"] for s in scripts]:
                        scripts.append({"kind": "make", "name": tgt,
                                        "cmd": f"make {tgt}"})
    if os.path.isfile(os.path.join(path, "manage.py")):
        scripts.append({"kind": "python", "name": "runserver",
                        "cmd": "python3 manage.py runserver"})
    return {"scripts": scripts[:40]}


def project_run(path, kind, name):
    path = os.path.realpath(path)
    if not path.startswith(HOME) or not os.path.isdir(path):
        raise ValueError("path not in your home")
    avail = project_scripts(path)["scripts"]
    match = next((s for s in avail if s["kind"] == kind and s["name"] == name),
                 None)
    if not match:
        raise ValueError("unknown script for this project")
    cmd = "cd " + _sh_quote(path) + " && " + match["cmd"]
    return start_job(["sh", "-c", cmd],
                     f"{match['cmd']} — {os.path.basename(path)}")


# ---------------------------------------------------- docker (extended) ------


def docker_stats():
    r = _run(["docker", "stats", "--no-stream", "--format",
                        "{{json .}}"], capture_output=True, text=True,
                       timeout=15)
    rows = []
    for line in r.stdout.splitlines():
        try:
            c = json.loads(line)
            rows.append({"name": c.get("Name"), "cpu": c.get("CPUPerc"),
                         "mem": c.get("MemPerc"), "memuse": c.get("MemUsage"),
                         "net": c.get("NetIO"), "block": c.get("BlockIO")})
        except json.JSONDecodeError:
            continue
    return {"stats": rows}


def docker_compose_projects():
    r = _run(["docker", "ps", "-a", "--format", "{{json .}}"],
                       capture_output=True, text=True, timeout=10)
    projects = {}
    for line in r.stdout.splitlines():
        try:
            c = json.loads(line)
        except json.JSONDecodeError:
            continue
        labels = c.get("Labels", "")
        proj = None
        for kv in labels.split(","):
            if kv.startswith("com.docker.compose.project="):
                proj = kv.split("=", 1)[1]
        if not proj:
            continue
        p = projects.setdefault(proj, {"name": proj, "running": 0, "total": 0})
        p["total"] += 1
        if c.get("State") == "running":
            p["running"] += 1
    return {"projects": sorted(projects.values(), key=lambda x: x["name"])}


def docker_compose_action(project, action):
    if action not in ("up", "down", "restart") \
            or not re.fullmatch(r"[\w.-]+", project or ""):
        raise ValueError("bad compose action")
    sub = {"up": ["up", "-d"], "down": ["down"], "restart": ["restart"]}[action]
    return start_job(["docker", "compose", "-p", project, *sub],
                     f"compose {action} — {project}")


def docker_disk():
    """`docker system df` — what images/containers/volumes actually cost, so
    you can see what pruning would reclaim before you prune."""
    r = _run(["docker", "system", "df", "--format", "{{json .}}"],
             capture_output=True, text=True, timeout=20)
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or "docker unavailable")[:300])
    rows = []
    for line in r.stdout.splitlines():
        if not line.strip():
            continue
        try:
            d = json.loads(line)
        except ValueError:
            continue
        rows.append({"type": d.get("Type", ""), "total": d.get("TotalCount", ""),
                     "active": d.get("Active", ""), "size": d.get("Size", ""),
                     "reclaimable": d.get("Reclaimable", "")})
    volumes = []
    rv = _run(["docker", "volume", "ls", "--format", "{{json .}}"],
              capture_output=True, text=True, timeout=15)
    if rv.returncode == 0:
        for line in rv.stdout.splitlines():
            if not line.strip():
                continue
            try:
                d = json.loads(line)
            except ValueError:
                continue
            volumes.append({"name": d.get("Name", ""), "driver": d.get("Driver", ""),
                            "size": d.get("Size", "")})
    return {"usage": rows, "volumes": volumes[:60]}


def docker_prune(kind):
    if kind == "images":
        return start_job(["docker", "image", "prune", "-f"],
                         "Prune dangling images")
    if kind == "system":
        return start_job(["docker", "system", "prune", "-f"],
                         "Prune stopped containers, networks, dangling images")
    raise ValueError("bad prune kind")


# --------------------------------------------------------------- monitor -----

CFG_DIR = os.path.join(HOME, ".config/perch")
MON_CFG_FILE = os.path.join(CFG_DIR, "monitor.json")
MON_DIR = os.path.join(HOME, ".cache/perch")
ALERTS_FILE = os.path.join(MON_DIR, "alerts.jsonl")
MINUTES_FILE = os.path.join(MON_DIR, "history.jsonl")

MON_DEFAULTS = {
    "cpu": {"on": True, "th": 90, "label": "CPU above % (sustained 60 s)"},
    "mem": {"on": True, "th": 95, "label": "Memory above %"},
    "temp": {"on": True, "th": 85, "label": "Temperature above °C"},
    "disk": {"on": True, "th": 90, "label": "Any disk above % full"},
    "battery": {"on": True, "th": 15, "label": "Battery below % (unplugged)"},
}


def mon_cfg():
    cfg = {k: dict(v) for k, v in MON_DEFAULTS.items()}
    try:
        with open(MON_CFG_FILE) as f:
            saved = json.load(f)
        for k in cfg:
            if k in saved:
                cfg[k]["on"] = bool(saved[k].get("on", cfg[k]["on"]))
                cfg[k]["th"] = float(saved[k].get("th", cfg[k]["th"]))
    except (OSError, ValueError):
        pass
    return cfg


def mon_save(new):
    os.makedirs(CFG_DIR, exist_ok=True)
    cfg = mon_cfg()
    for k in cfg:
        if k in new:
            cfg[k]["on"] = bool(new[k].get("on"))
            cfg[k]["th"] = float(new[k].get("th", cfg[k]["th"]))
    atomic_write(MON_CFG_FILE, json.dumps(cfg))
    return cfg


# ---- custom rules: watch a unit, a port, a process or a folder's size ----
# The five built-in rules above cover machine-wide metrics; these cover the
# specific things a given machine is supposed to be doing. They ride the same
# firing path (cooldown, master switch, channels, history) as everything else.

CUSTOM_FILE = os.path.join(CFG_DIR, "customrules.json")
CUSTOM_MAX = 24
CUSTOM_KINDS = ("unit", "port", "process", "path")
CUSTOM_LABELS = {
    "unit": "systemd user unit is not running",
    "port": "nothing is listening on port",
    "process": "no running process matches",
    "path": "folder is larger than (GB)",
}


def _custom_clean(r):
    """Validate one rule, or raise. Returns the stored form."""
    kind = r.get("kind")
    if kind not in CUSTOM_KINDS:
        raise ValueError("unknown rule type")
    target = str(r.get("target", "")).strip()
    if not target:
        raise ValueError("rule needs something to watch")
    if kind == "unit":
        if not SVC_RE.fullmatch(target):
            raise ValueError(f"'{target}' is not a unit name (try foo.service)")
    elif kind == "port":
        if not target.isdigit() or not 1 <= int(target) <= 65535:
            raise ValueError(f"'{target}' is not a port number")
        target = str(int(target))
    elif kind == "process":
        if len(target) > 64:
            raise ValueError("process pattern is too long")
    elif kind == "path":
        target = os.path.realpath(os.path.expanduser(target))
    value = r.get("value", 0)
    try:
        value = float(value or 0)
    except (TypeError, ValueError):
        raise ValueError("threshold must be a number")
    name = str(r.get("name", "")).strip()[:60] or f"{kind}:{target}"
    return {"name": name, "kind": kind, "target": target, "value": value,
            "enabled": bool(r.get("enabled", True))}


def custom_rules():
    try:
        with open(CUSTOM_FILE) as f:
            saved = json.load(f)
    except (OSError, ValueError):
        return []
    out = []
    for r in (saved if isinstance(saved, list) else [])[:CUSTOM_MAX]:
        try:
            out.append(_custom_clean(r))
        except ValueError:
            continue          # drop anything that no longer validates
    return out


def custom_save(rules):
    if not isinstance(rules, list):
        raise ValueError("expected a list of rules")
    if len(rules) > CUSTOM_MAX:
        raise ValueError(f"at most {CUSTOM_MAX} custom rules")
    cleaned = [_custom_clean(r) for r in rules]
    os.makedirs(CFG_DIR, exist_ok=True)
    atomic_write(CUSTOM_FILE, json.dumps(cleaned))
    return cleaned


def _unit_active(unit):
    r = _run(["systemctl", "--user", "is-active", unit],
             capture_output=True, text=True, timeout=8)
    return r.stdout.strip() == "active"


def _port_listening(port):
    try:
        for c in psutil.net_connections(kind="inet"):
            if c.status == "LISTEN" and c.laddr.port == port:
                return True
    except psutil.Error:
        return True           # can't tell — don't cry wolf
    return False


def _process_running(pattern):
    pat = pattern.lower()
    for p in psutil.process_iter(["name"]):
        try:
            if pat in (p.info["name"] or "").lower():
                return True
        except psutil.Error:
            continue
    return False


def fmt_bytes(n):
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if n < 1024 or unit == "TB":
            return f"{n:.0f} {unit}" if n >= 10 or unit == "B" else f"{n:.1f} {unit}"
        n /= 1024.0


def custom_check(rule):
    """Return an alert message when the rule is breached, else None."""
    kind, target = rule["kind"], rule["target"]
    if kind == "unit":
        if not _unit_active(target):
            return f"{rule['name']} — {target} is not running"
    elif kind == "port":
        if not _port_listening(int(target)):
            return f"{rule['name']} — nothing is listening on port {target}"
    elif kind == "process":
        if not _process_running(target):
            return f"{rule['name']} — no process matching '{target}' is running"
    elif kind == "path":
        limit = rule["value"]
        if limit <= 0:
            return None
        size = quick_size(target, budget=4)
        if size > limit * (1024 ** 3):
            return (f"{rule['name']} — {target} is {fmt_bytes(size)}, "
                    f"over the {limit:g} GB limit")
    return None


def custom_tick():
    """Evaluate every enabled custom rule; called once a minute."""
    for i, rule in enumerate(custom_rules()):
        if not rule.get("enabled", True):
            continue
        try:
            msg = custom_check(rule)
        except Exception:  # noqa: BLE001 — one bad rule must not stop the rest
            continue
        if msg:
            _fire(f"custom:{i}:{rule['kind']}:{rule['target']}", msg, None)


# ---- alert firing state: breach tracking, per-rule cooldown, master switch ---

_mon_breach = {}    # rule -> breach start ts
_mon_fired = {}     # rule -> last fired ts
_mon_minute = 0.0
COOLDOWN = 600

ALERTCTL_FILE = os.path.join(CFG_DIR, "alertctl.json")
SNOOZE_MAX = 7 * 24 * 60          # minutes — a week is plenty for "shut up"


def _alert_ctl_write(st):
    os.makedirs(CFG_DIR, exist_ok=True)
    atomic_write(ALERTCTL_FILE, json.dumps(st))
    return st


def alert_ctl():
    """Current alerting state. An expired snooze resumes alerting by itself."""
    st = {"enabled": True, "until": 0.0}
    try:
        with open(ALERTCTL_FILE) as f:
            saved = json.load(f)
        st["enabled"] = bool(saved.get("enabled", True))
        st["until"] = float(saved.get("until", 0) or 0)
    except (OSError, ValueError, TypeError):
        return st
    if not st["enabled"] and st["until"] and time.time() >= st["until"]:
        st = _alert_ctl_write({"enabled": True, "until": 0.0})
    return st


def alerts_paused():
    return not alert_ctl()["enabled"]


def alert_ctl_set(action, minutes=0):
    """start | stop | snooze (for `minutes`) | clear (wipe the history)."""
    if action == "start":
        return _alert_ctl_write({"enabled": True, "until": 0.0})
    if action == "stop":
        return _alert_ctl_write({"enabled": False, "until": 0.0})
    if action == "snooze":
        mins = max(1, min(SNOOZE_MAX, int(minutes or 60)))
        return _alert_ctl_write({"enabled": False,
                                 "until": time.time() + mins * 60})
    if action == "clear":
        try:
            os.remove(ALERTS_FILE)
        except OSError:
            pass
        _mon_fired.clear()
        _mon_breach.clear()
        return alert_ctl()
    raise ValueError("unknown alert action")


# ---- home-screen layout, stored server-side so it follows the user ----
# The browser keeps a localStorage copy as an offline cache; this file is the
# source of truth, which is what makes the layout survive a different browser.

HOME_LAYOUT_FILE = os.path.join(CFG_DIR, "home.json")


def home_layout():
    try:
        with open(HOME_LAYOUT_FILE) as f:
            saved = json.load(f)
    except (OSError, ValueError):
        return {"layout": None}
    return {"layout": saved if isinstance(saved, dict) else None}


def home_layout_save(body):
    """Store {order, hidden, sizes}. Ids are opaque to the server — the widget
    catalogue lives in the frontend — so we only bound shape and size."""
    layout = body.get("layout") if isinstance(body, dict) else None
    if layout is None:
        try:
            os.remove(HOME_LAYOUT_FILE)
        except OSError:
            pass
        return {"layout": None}
    if not isinstance(layout, dict):
        raise ValueError("layout must be an object")
    ident = re.compile(r"[\w-]{1,40}")
    order = [w for w in (layout.get("order") or [])[:120]
             if isinstance(w, str) and ident.fullmatch(w)]
    hidden = [w for w in (layout.get("hidden") or [])[:120]
              if isinstance(w, str) and ident.fullmatch(w)]
    sizes = {k: v for k, v in (layout.get("sizes") or {}).items()
             if isinstance(k, str) and ident.fullmatch(k)
             and v in ("s", "m", "l", "full")}
    clean = {"order": order, "hidden": hidden, "sizes": sizes}
    os.makedirs(CFG_DIR, exist_ok=True)
    atomic_write(HOME_LAYOUT_FILE, json.dumps(clean))
    return {"layout": clean}


def notify(title, msg, critical=True):
    try:
        subprocess.Popen(["notify-send", "--app-name=Perch",
                          "--icon=perch",
                          "--urgency=" + ("critical" if critical else "normal"),
                          title, msg],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except OSError:
        pass


# ---- outbound notification channels (ntfy / Slack / Discord / webhook) ----

def _notify_file():
    return os.path.join(CFG_DIR, "notify.json")


def notify_cfg():
    cfg = {"desktop": True, "channels": []}
    try:
        with open(_notify_file()) as f:
            cfg.update(json.load(f))
    except (OSError, ValueError):
        pass
    return cfg


def notify_save(body):
    os.makedirs(CFG_DIR, exist_ok=True)
    cfg = notify_cfg()
    if "desktop" in body:
        cfg["desktop"] = bool(body["desktop"])
    if "channels" in body:
        cfg["channels"] = [c for c in body["channels"]
                           if c.get("type") and c.get("url")][:12]
    atomic_write(_notify_file(), json.dumps(cfg))
    os.chmod(_notify_file(), 0o600)
    return cfg


def _post_channel(ch, title, msg):
    import urllib.request as ur
    url, t = ch["url"], ch["type"]
    try:
        if t == "ntfy":
            req = ur.Request(url, data=msg.encode(), method="POST",
                             headers={"Title": title, "Priority": "high",
                                      "Tags": "warning"})
        elif t == "slack":
            req = ur.Request(url, method="POST",
                             headers={"Content-Type": "application/json"},
                             data=json.dumps({"text": f"*{title}*\n{msg}"}).encode())
        elif t == "discord":
            req = ur.Request(url, method="POST",
                             headers={"Content-Type": "application/json"},
                             data=json.dumps({"content": f"**{title}**\n{msg}"}).encode())
        else:  # generic webhook
            req = ur.Request(url, method="POST",
                             headers={"Content-Type": "application/json"},
                             data=json.dumps({"title": title, "message": msg,
                                              "host": os.uname().nodename,
                                              "time": time.time()}).encode())
        ur.urlopen(req, timeout=10).read()
        return True
    except Exception:  # noqa: BLE001 — a broken channel must not break alerting
        return False


def dispatch(title, msg, critical=True):
    """Fan a notification out to the desktop and every enabled channel."""
    cfg = notify_cfg()
    if cfg.get("desktop", True):
        notify(title, msg, critical)
    for ch in cfg.get("channels", []):
        if ch.get("enabled", True):
            threading.Thread(target=_post_channel, args=(ch, title, msg),
                             daemon=True).start()


def notify_test():
    dispatch("Perch test", "If you can read this, alerts are wired up 🎉",
             critical=False)
    n = sum(1 for c in notify_cfg().get("channels", []) if c.get("enabled", True))
    return {"ok": True, "channels": n}


def _alert(rule, msg, value):
    """Record and send one alert. Silent while alerting is stopped/snoozed —
    nothing is logged either, so the history stays a record of what was sent."""
    if alerts_paused():
        return
    os.makedirs(MON_DIR, exist_ok=True)
    with open(ALERTS_FILE, "a") as f:
        f.write(json.dumps({"t": time.time(), "rule": rule, "msg": msg,
                            "value": value}) + "\n")
    dispatch("⚠ " + msg.split(" — ")[0], msg)


def _fire(rule, msg, value):
    # checked before the cooldown is stamped: a rule that breached while
    # alerting was stopped must be free to alert the moment it's started again
    if alerts_paused():
        return
    now = time.time()
    if now - _mon_fired.get(rule, 0) < COOLDOWN:
        return
    _mon_fired[rule] = now
    _alert(rule, msg, value)


def monitor_tick(entry):
    global _mon_minute
    cfg = mon_cfg()
    now = entry["t"]
    if cfg["cpu"]["on"]:
        if entry["cpu"] >= cfg["cpu"]["th"]:
            _mon_breach.setdefault("cpu", now)
            if now - _mon_breach["cpu"] >= 60:
                _fire("cpu", f"CPU at {entry['cpu']:.0f}% for over a minute "
                             f"— threshold {cfg['cpu']['th']:.0f}%",
                      entry["cpu"])
        else:
            _mon_breach.pop("cpu", None)
    if cfg["mem"]["on"] and entry["mem"] >= cfg["mem"]["th"]:
        _fire("mem", f"Memory at {entry['mem']:.0f}% — threshold "
                     f"{cfg['mem']['th']:.0f}%", entry["mem"])
    t = entry.get("temp")
    if cfg["temp"]["on"] and t and t >= cfg["temp"]["th"]:
        _fire("temp", f"Temperature {t:.0f}°C — threshold "
                      f"{cfg['temp']['th']:.0f}°C", t)
    if now - _mon_minute >= 60:
        _mon_minute = now
        batt = None
        try:
            b = psutil.sensors_battery()
            if b:
                batt = round(b.percent)
                if (cfg["battery"]["on"] and not b.power_plugged
                        and b.percent <= cfg["battery"]["th"]):
                    _fire("battery", f"Battery at {b.percent:.0f}% and "
                                     "discharging — plug in", b.percent)
        except Exception:  # noqa: BLE001
            pass
        worst = 0.0
        for d in psutil.disk_partitions(all=False):
            if d.fstype in ("squashfs", "tmpfs", "devtmpfs"):
                continue
            try:
                pct = psutil.disk_usage(d.mountpoint).percent
            except OSError:
                continue
            if pct > worst:
                worst = pct
            if cfg["disk"]["on"] and pct >= cfg["disk"]["th"]:
                _fire("disk", f"Disk {d.mountpoint} is {pct:.0f}% full "
                              f"— threshold {cfg['disk']['th']:.0f}%", pct)
        os.makedirs(MON_DIR, exist_ok=True)
        with open(MINUTES_FILE, "a") as f:
            f.write(json.dumps({"t": now, "cpu": round(entry["cpu"], 1),
                                "mem": round(entry["mem"], 1),
                                "temp": t, "batt": batt,
                                "disk": round(worst, 1)}) + "\n")
        _prune_jsonl(MINUTES_FILE, 4000)
        _prune_jsonl(ALERTS_FILE, 1000)
        # custom rules shell out (systemctl, folder sizing), so once a minute
        # on the sampler thread rather than on every 2 s tick
        if not alerts_paused():
            custom_tick()


def _prune_jsonl(path, keep):
    try:
        with open(path) as f:
            lines = f.readlines()
        if len(lines) > keep * 1.25:
            with open(path, "w") as f:
                f.writelines(lines[-keep:])
    except OSError:
        pass


def _tail_jsonl(path, n):
    try:
        with open(path) as f:
            return [json.loads(x) for x in f.readlines()[-n:]]
    except (OSError, ValueError):
        return []


def monitor_state(brief=False):
    """Rules, alerting state and recent alerts. `brief` drops the 24 h history
    (1440 samples) for callers that only want the events — e.g. home widgets."""
    st = {"cfg": mon_cfg(), "ctl": alert_ctl(), "custom": custom_rules(),
          "custom_kinds": CUSTOM_LABELS,
          "events": list(reversed(_tail_jsonl(ALERTS_FILE, 100)))}
    if not brief:
        st["history"] = _tail_jsonl(MINUTES_FILE, 1440)
    return st


# ------------------------------------------------------- log-pattern watch ---
# Poll the journal (system/user/kernel) or a file every LOGWATCH_INTERVAL s;
# lines matching a rule's regex fire a notification (deduped, with cooldown).

LOGWATCH_INTERVAL = 20
_logwatch_stop = threading.Event()
_logwatch_thread = [None]
_lw_fired = {}          # match-key -> last fired ts
_lw_offset = {}         # file path -> byte offset
_lw_started_at = [0.0]


def _logwatch_file():
    return os.path.join(CFG_DIR, "logwatch.json")


def logwatch_cfg():
    cfg = {"enabled": False, "rules": []}
    try:
        with open(_logwatch_file()) as f:
            cfg.update(json.load(f))
    except (OSError, ValueError):
        pass
    return cfg


def logwatch_save(body):
    os.makedirs(CFG_DIR, exist_ok=True)
    cfg = logwatch_cfg()
    if "enabled" in body:
        cfg["enabled"] = bool(body["enabled"])
    if "rules" in body:
        clean = []
        for r in body["rules"][:20]:
            if not r.get("pattern") or not r.get("name"):
                continue
            try:
                re.compile(r["pattern"])
            except re.error:
                continue
            clean.append({"name": r["name"], "pattern": r["pattern"],
                          "source": r.get("source", "system"),
                          "path": r.get("path", ""),
                          "enabled": bool(r.get("enabled", True))})
        cfg["rules"] = clean
    atomic_write(_logwatch_file(), json.dumps(cfg))
    if cfg["enabled"]:
        _logwatch_start()
    else:
        _logwatch_stop.set()
    return cfg


def _lw_journal_lines(source, since_iso):
    flags = {"system": ["--system"], "user": ["--user"],
             "kernel": ["--system", "-k"]}.get(source, ["--system"])
    r = _run(["journalctl", *flags, "-o", "cat", "--no-pager",
                        "--since", since_iso], capture_output=True, text=True,
                       timeout=15)
    return r.stdout.splitlines()


def _lw_file_lines(path):
    try:
        size = os.path.getsize(path)
    except OSError:
        return []
    off = _lw_offset.get(path)
    if off is None or off > size:          # first sight or truncation/rotation
        _lw_offset[path] = size
        return []
    with open(path, errors="replace") as f:
        f.seek(off)
        data = f.read()
        _lw_offset[path] = f.tell()
    return data.splitlines()


def _logwatch_loop():
    import datetime
    last = time.time()
    while not _logwatch_stop.wait(LOGWATCH_INTERVAL):
        cfg = logwatch_cfg()
        if not cfg["enabled"]:
            return
        since = datetime.datetime.fromtimestamp(last).strftime("%Y-%m-%d %H:%M:%S")
        last = time.time()
        for rule in cfg["rules"]:
            if not rule.get("enabled", True):
                continue
            try:
                pat = re.compile(rule["pattern"])
            except re.error:
                continue
            if rule.get("source") == "file" and rule.get("path"):
                lines = _lw_file_lines(os.path.realpath(rule["path"]))
            else:
                lines = _lw_journal_lines(rule.get("source", "system"), since)
            for line in lines:
                if not pat.search(line):
                    continue
                key = rule["name"] + "|" + line[:200]
                now = time.time()
                if now - _lw_fired.get(key, 0) < 300:
                    continue
                _lw_fired[key] = now
                _alert("logwatch:" + rule["name"],
                       f"{rule['name']} — {line[:300]}", None)
            if len(_lw_fired) > 500:
                _lw_fired.clear()


def _logwatch_start():
    _logwatch_stop.clear()
    _lw_offset.clear()  # re-seek files to current end so we don't replay
    # seed each file rule's offset now, so appends after start are caught on
    # the first poll (no startup swallow window)
    for rule in logwatch_cfg()["rules"]:
        if rule.get("source") == "file" and rule.get("path"):
            p = os.path.realpath(rule["path"])
            try:
                _lw_offset[p] = os.path.getsize(p)
            except OSError:
                pass
    if not (_logwatch_thread[0] and _logwatch_thread[0].is_alive()):
        _logwatch_thread[0] = threading.Thread(target=_logwatch_loop,
                                               daemon=True)
        _logwatch_thread[0].start()


def logwatch_state():
    c = logwatch_cfg()
    c["running"] = bool(_logwatch_thread[0] and _logwatch_thread[0].is_alive())
    return c


if logwatch_cfg()["enabled"]:
    _logwatch_start()


# ------------------------------------------------------------- http tester ---


def http_request(method, url, headers_text="", body="", timeout=20):
    import urllib.request as ur
    if not re.match(r"^https?://", url or ""):
        raise ValueError("URL must start with http:// or https://")
    hdrs = {}
    for line in (headers_text or "").splitlines():
        if ":" in line:
            k, v = line.split(":", 1)
            hdrs[k.strip()] = v.strip()
    method = (method or "GET").upper()
    data = body.encode() if body and method not in ("GET", "HEAD") else None
    req = ur.Request(url, method=method, data=data)
    for k, v in hdrs.items():
        req.add_header(k, v)
    t0 = time.time()
    try:
        with ur.urlopen(req, timeout=min(int(timeout or 20), 60)) as resp:
            raw = resp.read(512 * 1024)
            status, reason = resp.status, resp.reason
            rheaders = dict(resp.headers)
    except urllib.error.HTTPError as e:
        raw = e.read(512 * 1024)
        status, reason, rheaders = e.code, e.reason, dict(e.headers)
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"request failed: {e}")
    ms = round((time.time() - t0) * 1000)
    curl = ["curl", "-X", method, f"'{url}'"]
    for k, v in hdrs.items():
        curl.append(f"-H '{k}: {v}'")
    if data:
        curl.append(f"-d '{body}'")
    return {"status": status, "reason": reason, "ms": ms,
            "headers": rheaders, "size": len(raw),
            "body": raw.decode("utf-8", errors="replace"),
            "curl": " ".join(curl)}


# ------------------------------------------------- updates / net extras ------

_upd_cache = {"t": 0.0, "data": None}


def _parse_updates(pm, out):
    pkgs = []
    if pm == "apt":
        for line in out.splitlines():
            m = re.match(r"^([^/]+)/(\S+)\s+(\S+)\s+\S+\s+\[upgradable from: "
                         r"(.+)\]", line)
            if m:
                pkgs.append({"name": m.group(1), "repo": m.group(2),
                             "new": m.group(3), "old": m.group(4),
                             "security": "-security" in m.group(2)})
    elif pm == "dnf":
        for line in out.splitlines():
            m = re.match(r"^(\S+)\.\S+\s+(\S+)\s+(\S+)\s*$", line)
            if m:
                pkgs.append({"name": m.group(1), "repo": m.group(3),
                             "new": m.group(2), "old": "",
                             "security": False})
    elif pm == "pacman":
        for line in out.splitlines():
            m = re.match(r"^(\S+)\s+(\S+)\s+->\s+(\S+)", line)
            if m:
                pkgs.append({"name": m.group(1), "repo": "",
                             "new": m.group(3), "old": m.group(2),
                             "security": False})
    elif pm == "zypper":
        for line in out.splitlines():
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 5 and parts[0] == "v":
                pkgs.append({"name": parts[2], "repo": parts[1],
                             "new": parts[4], "old": parts[3],
                             "security": False})
    return pkgs


def pkg_updates(force=False):
    if not force and _upd_cache["data"] and time.time() - _upd_cache["t"] < 600:
        return _upd_cache["data"]
    cmds = {"apt": ["apt", "list", "--upgradable"],
            "dnf": ["dnf", "-q", "check-update"],
            "pacman": ["pacman", "-Qu"],
            "zypper": ["zypper", "-q", "lu"]}
    pkgs = []
    if _PM:
        r = _run(cmds[_PM], capture_output=True, text=True, timeout=60,
                 env={**os.environ, "LC_ALL": "C"})
        pkgs = _parse_updates(_PM, r.stdout)
    st = 0.0
    try:
        st = os.path.getmtime("/var/lib/apt/lists")
    except OSError:
        pass
    data = {"packages": pkgs[:300], "count": len(pkgs), "pm": _PM,
            "security": sum(1 for p in pkgs if p["security"]),
            "lists_updated": st}
    _upd_cache.update(t=time.time(), data=data)
    return data


_pubip_cache = {"t": 0.0, "ip": None}


UA = {"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) perch"}


def _get(url, timeout):
    import urllib.request as ur
    return ur.urlopen(ur.Request(url, headers=UA), timeout=timeout)


def public_ip():
    if _pubip_cache["ip"] and time.time() - _pubip_cache["t"] < 900:
        return _pubip_cache["ip"]
    try:
        ip = _get("https://api.ipify.org", 6).read().decode()
        _pubip_cache.update(t=time.time(), ip=ip)
        return ip
    except Exception:  # noqa: BLE001
        return None


def speed_test():
    pings = []
    for _ in range(3):
        t0 = time.time()
        try:
            _get("https://speed.cloudflare.com/__down?bytes=0", 8).read()
            pings.append((time.time() - t0) * 1000)
        except Exception:  # noqa: BLE001
            pass
    nbytes = 20_000_000
    t0 = time.time()
    try:
        with _get(f"https://speed.cloudflare.com/__down?bytes={nbytes}",
                  45) as resp:
            got = 0
            while chunk := resp.read(256 * 1024):
                got += len(chunk)
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"speed test failed: {e}")
    secs = time.time() - t0
    return {"mbps": round(got * 8 / secs / 1e6, 1),
            "ping_ms": round(min(pings), 1) if pings else None,
            "mb": round(got / 1e6, 1), "secs": round(secs, 1),
            "ip": public_ip()}


def proc_detail(pid):
    p = psutil.Process(pid)
    with p.oneshot():
        try:
            conns = len(p.net_connections())
        except (psutil.AccessDenied, AttributeError):
            try:
                conns = len(p.connections())
            except Exception:  # noqa: BLE001
                conns = None
        mem = p.memory_info()
        parent = p.parent()
        d = {
            "pid": p.pid, "name": p.name(), "status": p.status(),
            "user": p.username(), "started": p.create_time(),
            "cmdline": " ".join(p.cmdline())[:800],
            "exe": (p.exe() if os.access(f"/proc/{pid}/exe", os.R_OK)
                    else None),
            "cwd": None, "rss": mem.rss, "vms": mem.vms,
            "threads": p.num_threads(), "fds": None,
            "conns": conns, "nice": p.nice(),
            "cpu": p.cpu_percent(interval=0.15),
            "parent": f"{parent.name()} ({parent.pid})" if parent else None,
            "children": [f"{c.name()} ({c.pid})"
                         for c in p.children()[:10]],
        }
        for field, fn in (("cwd", p.cwd), ("fds", p.num_fds)):
            try:
                d[field] = fn()
            except psutil.Error:
                pass
    return d


# ------------------------------------------------------- caps / raw / yaml ---


def capabilities():
    editor = next((e for e in ("code", "subl", "gedit") if shutil.which(e)),
                  None)
    return {
        "editor": editor,
        "terminal": bool(shutil.which("gnome-terminal")
                         or shutil.which("x-terminal-emulator")),
        "opener": bool(shutil.which("xdg-open")),
        "docker": bool(shutil.which("docker")),
        "ai": os.path.exists(CLAUDE_BIN),
        "notify": bool(shutil.which("notify-send")),
        "nmcli": bool(shutil.which("nmcli")),
        "yaml": True,
        "native_pm": _PM,
        "snap": _HAS_SNAP,
        "flatpak": _HAS_FLATPAK,
        # a working GNOME gsettings (needs the session dbus, not just the CLI)
        "gnome": bool(shutil.which("gsettings"))
                 and _gset(_IFACE, "color-scheme") is not None,
        "battery": _has_battery(),
        "wayland": os.environ.get("XDG_SESSION_TYPE") == "wayland",
    }


def _has_battery():
    try:
        return psutil.sensors_battery() is not None
    except (OSError, RuntimeError):
        return False


MIMES = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
         ".gif": "image/gif", ".webp": "image/webp", ".svg": "image/svg+xml",
         ".ico": "image/x-icon", ".bmp": "image/bmp",
         ".pdf": "application/pdf", ".mp4": "video/mp4",
         ".webm": "video/webm", ".mkv": "video/x-matroska",
         ".mp3": "audio/mpeg", ".wav": "audio/wav", ".ogg": "audio/ogg",
         ".m4a": "audio/mp4", ".flac": "audio/flac",
         ".html": "text/html", ".htm": "text/html"}
PREVIEWABLE = {"image": (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg",
                         ".ico", ".bmp"),
               "pdf": (".pdf",),
               "video": (".mp4", ".webm", ".mkv"),
               "audio": (".mp3", ".wav", ".ogg", ".m4a", ".flac")}


def preview_kind(name):
    ext = os.path.splitext(name.lower())[1]
    for kind, exts in PREVIEWABLE.items():
        if ext in exts:
            return kind
    return None


def yaml_convert(text, direction):
    import yaml
    if direction == "y2j":
        return json.dumps(yaml.safe_load(text), indent=2, default=str)
    return yaml.safe_dump(json.loads(text), sort_keys=False,
                          allow_unicode=True)


# ---------------------------------------------------------- job runner -------

JOBS = {}          # id -> {"cmd","lines","status","code","started","title"}
_job_seq = [0]
_job_lock = threading.Lock()


def start_job(argv, title, privileged=False):
    with _job_lock:
        _job_seq[0] += 1
        jid = str(_job_seq[0])
    if privileged and os.geteuid() != 0:
        # the wrapper gives a branded polkit prompt + session-long auth cache
        wrapper = shutil.which("perch-pkexec")
        argv = ["pkexec", wrapper, *argv] if wrapper else ["pkexec", *argv]
    job = {"id": jid, "title": title, "cmd": " ".join(argv),
           "lines": ["$ " + " ".join(argv), ""], "status": "running",
           "code": None, "started": time.time()}
    JOBS[jid] = job
    # keep only the 30 most recent jobs
    if len(JOBS) > 30:
        for k in sorted(JOBS, key=lambda k: JOBS[k]["started"])[:-30]:
            JOBS.pop(k, None)

    def run():
        try:
            p = subprocess.Popen(argv, stdout=subprocess.PIPE,
                                 stderr=subprocess.STDOUT, text=True,
                                 bufsize=1,
                                 env={**os.environ, "DEBIAN_FRONTEND":
                                      "noninteractive"})
            for line in p.stdout:
                job["lines"].append(line.rstrip("\n")[:400])
                if len(job["lines"]) > 1000:
                    job["lines"] = job["lines"][-1000:]
            p.wait()
            job["code"] = p.returncode
            job["status"] = "done" if p.returncode == 0 else "failed"
            if p.returncode in (126, 127):
                job["lines"].append("(authentication cancelled or failed)")
        except Exception as e:  # noqa: BLE001
            job["lines"].append("ERROR: " + str(e))
            job["status"] = "failed"
            job["code"] = -1

    threading.Thread(target=run, daemon=True).start()
    return {"id": jid}


def job_status(jid, since=0):
    job = JOBS.get(jid)
    if not job:
        raise ValueError("unknown job")
    since = int(since or 0)
    return {"id": jid, "status": job["status"], "code": job["code"],
            "title": job["title"],
            "lines": job["lines"][since:], "total": len(job["lines"])}


# ------------------------------------------------------------- packages ------
# Native package manager abstraction: apt / dnf / pacman / zypper detected at
# startup, plus snap and flatpak as universal extras when installed.


def _native_pm():
    for pm in ("apt", "dnf", "pacman", "zypper"):
        if shutil.which(pm):
            return pm
    return None


_PM = _native_pm()
_HAS_SNAP = bool(shutil.which("snap"))
_HAS_FLATPAK = bool(shutil.which("flatpak"))

_PM_SEARCH = {
    "apt": lambda q: ["apt-cache", "search", "--names-only", q],
    "dnf": lambda q: ["dnf", "-q", "search", q],
    "pacman": lambda q: ["pacman", "-Ss", q],
    "zypper": lambda q: ["zypper", "-q", "se", q],
}
_PM_INSTALL = {
    "apt": lambda n: ["apt-get", "install", "-y", n],
    "dnf": lambda n: ["dnf", "install", "-y", n],
    "pacman": lambda n: ["pacman", "-S", "--noconfirm", n],
    "zypper": lambda n: ["zypper", "-n", "install", n],
}
_PM_REMOVE = {
    "apt": lambda n: ["apt-get", "remove", "-y", n],
    "dnf": lambda n: ["dnf", "remove", "-y", n],
    "pacman": lambda n: ["pacman", "-R", "--noconfirm", n],
    "zypper": lambda n: ["zypper", "-n", "remove", n],
}
_PM_UPGRADE = {
    "apt": ["sh", "-c", "apt-get update && apt-get upgrade -y"],
    "dnf": ["dnf", "upgrade", "-y"],
    "pacman": ["pacman", "-Syu", "--noconfirm"],
    "zypper": ["sh", "-c", "zypper -n refresh && zypper -n update"],
}


def _parse_pm_search(pm, out):
    pkgs = []
    if pm == "apt":
        for line in out.splitlines():
            if " - " in line:
                name, desc = line.split(" - ", 1)
                pkgs.append({"name": name.strip(), "desc": desc[:120]})
    elif pm == "dnf":
        for line in out.splitlines():
            m = re.match(r"^(\S+)\.\S+\s*:\s*(.*)", line)
            if m:
                pkgs.append({"name": m.group(1), "desc": m.group(2)[:120]})
    elif pm == "pacman":
        cur = None
        for line in out.splitlines():
            m = re.match(r"^\S+/(\S+)\s+\S+", line)
            if m:
                cur = {"name": m.group(1), "desc": ""}
                pkgs.append(cur)
            elif cur is not None and line[:1] in (" ", "\t"):
                cur["desc"] = (cur["desc"] + " " + line.strip()).strip()[:120]
    elif pm == "zypper":
        for line in out.splitlines():
            parts = [p.strip() for p in line.split("|")]
            if (len(parts) >= 4 and parts[1] and parts[1] != "Name"
                    and not set(parts[1]) <= {"-", "+"}):
                pkgs.append({"name": parts[1], "desc": parts[2][:120]})
    return pkgs[:40]


def _installed_native():
    if _PM == "apt":
        r = _run(["dpkg-query", "-f", "${Package}\n", "-W"],
                 capture_output=True, text=True)
    elif _PM in ("dnf", "zypper"):
        r = _run(["rpm", "-qa", "--qf", "%{NAME}\n"],
                 capture_output=True, text=True, timeout=30)
    elif _PM == "pacman":
        r = _run(["pacman", "-Qq"], capture_output=True, text=True)
    else:
        return set()
    return set(r.stdout.split())


def pkg_search(q):
    q = q.strip()
    out = {"native_pm": _PM, "native": [], "snap": [], "flatpak": []}
    if len(q) < 2:
        return out
    if _PM:
        r = _run(_PM_SEARCH[_PM](q), capture_output=True, text=True,
                 timeout=30, env={**os.environ, "LC_ALL": "C"})
        installed = _installed_native()
        out["native"] = [{**p, "installed": p["name"] in installed}
                         for p in _parse_pm_search(_PM, r.stdout)]
    if _HAS_SNAP:
        rs = _run(["snap", "find", q], capture_output=True, text=True,
                  timeout=25, env={**os.environ, "LC_ALL": "C"})
        rsl = _run(["snap", "list"], capture_output=True, text=True)
        snap_inst = {ln.split()[0] for ln in rsl.stdout.splitlines()[1:] if ln}
        for line in rs.stdout.splitlines()[1:21]:
            parts = line.split(None, 4)
            if len(parts) >= 5:
                out["snap"].append(
                    {"name": parts[0], "version": parts[1],
                     "publisher": parts[2], "desc": parts[4][:120],
                     "installed": parts[0] in snap_inst})
    if _HAS_FLATPAK:
        rf = _run(["flatpak", "search",
                   "--columns=application,name,description", q],
                  capture_output=True, text=True, timeout=30)
        ri = _run(["flatpak", "list", "--columns=application"],
                  capture_output=True, text=True)
        flat_inst = set(ri.stdout.split())
        for line in rf.stdout.splitlines()[:20]:
            parts = line.split("\t")
            if len(parts) >= 2 and "." in parts[0]:
                out["flatpak"].append(
                    {"name": parts[0], "title": parts[1],
                     "desc": (parts[2] if len(parts) > 2 else "")[:120],
                     "installed": parts[0] in flat_inst})
    return out


def pkg_install(mgr, name):
    if not re.fullmatch(r"[a-zA-Z0-9][a-zA-Z0-9.+_-]{0,80}", name):
        raise ValueError("invalid package name")
    if mgr in ("native", "apt") and _PM:
        return start_job(_PM_INSTALL[_PM](name),
                         f"Install {name} ({_PM})", privileged=True)
    if mgr in ("native-remove", "apt-remove") and _PM:
        return start_job(_PM_REMOVE[_PM](name),
                         f"Remove {name} ({_PM})", privileged=True)
    if mgr == "snap":
        return start_job(["snap", "install", name],
                         f"Install {name} (snap)", privileged=True)
    if mgr == "snap-remove":
        return start_job(["snap", "remove", name],
                         f"Remove {name} (snap)", privileged=True)
    if mgr == "flatpak":  # flatpak talks to polkit itself — no pkexec
        return start_job(["flatpak", "install", "-y", "--noninteractive",
                          name], f"Install {name} (flatpak)")
    if mgr == "flatpak-remove":
        return start_job(["flatpak", "uninstall", "-y", name],
                         f"Remove {name} (flatpak)")
    raise ValueError("unknown package manager")


def upgrade_all(mgr):
    _upd_cache["t"] = 0
    if mgr in ("native", "apt") and _PM:
        return start_job(_PM_UPGRADE[_PM],
                         f"Upgrade all {_PM} packages", privileged=True)
    if mgr == "snap":
        return start_job(["snap", "refresh"], "Refresh all snaps",
                         privileged=True)
    if mgr == "flatpak":
        return start_job(["flatpak", "update", "-y"],
                         "Update all flatpaks")
    raise ValueError("unknown package manager")


# -------------------------------------------------------------- settings -----


def _gset(schema, key):
    r = _run(["gsettings", "get", schema, key],
                       capture_output=True, text=True)
    return r.stdout.strip().strip("'") if r.returncode == 0 else None


def _gbool(schema, key):
    v = _gset(schema, key)
    return None if v is None else v == "true"


def _gnum(schema, key):
    v = _gset(schema, key)
    if v is None:
        return None
    v = re.sub(r"^\s*u?int\d+\s+", "", v)  # drop a "uint32 "/"int64 " type tag
    m = re.search(r"-?\d+(?:\.\d+)?", v)
    if not m:
        return None
    return float(m.group()) if "." in m.group() else int(m.group())


def _gset_write(schema, key, val):
    _run(["gsettings", "set", schema, key, val], capture_output=True)


_TOUCHPAD = "org.gnome.desktop.peripherals.touchpad"
_COLOR = "org.gnome.settings-daemon.plugins.color"
_IFACE = "org.gnome.desktop.interface"
_POWER = "org.gnome.settings-daemon.plugins.power"
_WM = "org.gnome.desktop.wm.preferences"
_MUTTER = "org.gnome.mutter"
_MOUSE = "org.gnome.desktop.peripherals.mouse"


def _tweak_options():
    """Installed GTK/icon/cursor themes and font families, for the Tweaks UI."""
    gtk = set()
    for root in ("/usr/share/themes", "~/.themes", "~/.local/share/themes"):
        try:
            for d in os.scandir(os.path.expanduser(root)):
                if d.is_dir():
                    gtk.add(d.name)
        except OSError:
            pass
    icons, cursors = set(), set()
    for root in ("/usr/share/icons", "~/.icons", "~/.local/share/icons"):
        try:
            for d in os.scandir(os.path.expanduser(root)):
                if not d.is_dir():
                    continue
                if os.path.isdir(os.path.join(d.path, "cursors")):
                    cursors.add(d.name)
                if os.path.isfile(os.path.join(d.path, "index.theme")):
                    icons.add(d.name)
        except OSError:
            pass
    fonts = set()
    r = _run(["fc-list", ":", "family"], capture_output=True,
                       text=True)
    if r.returncode == 0:
        for line in r.stdout.splitlines():
            fam = line.split(",")[0].strip()
            if fam:
                fonts.add(fam)
    return {"gtk_themes": sorted(gtk), "icon_themes": sorted(icons),
            "cursor_themes": sorted(cursors), "fonts": sorted(fonts)}


_TWEAK_BOOLS = {
    "animations": (_IFACE, "enable-animations"),
    "hot_corner": (_IFACE, "enable-hot-corners"),
    "clock_weekday": (_IFACE, "clock-show-weekday"),
    "clock_date": (_IFACE, "clock-show-date"),
    "clock_seconds": (_IFACE, "clock-show-seconds"),
    "ws_dynamic": (_MUTTER, "dynamic-workspaces"),
}
_TWEAK_CHOICES = {
    "font_aa": (_IFACE, "font-antialiasing",
                ("none", "grayscale", "rgba")),
    "font_hint": (_IFACE, "font-hinting",
                  ("none", "slight", "medium", "full")),
}
_TWEAK_THEMES = {
    "gtk_theme": (_IFACE, "gtk-theme", "gtk_themes"),
    "icon_theme": (_IFACE, "icon-theme", "icon_themes"),
    "cursor_theme": (_IFACE, "cursor-theme", "cursor_themes"),
}
_TWEAK_FONTS = {
    "font_name": "font-name",
    "mono_font": "monospace-font-name",
    "doc_font": "document-font-name",
}
_TITLEBAR_LAYOUTS = {
    "close": "appmenu:close",
    "max-close": "appmenu:maximize,close",
    "min-max-close": "appmenu:minimize,maximize,close",
}


def _power_profile():
    r = _run(["powerprofilesctl", "get"], capture_output=True,
                       text=True)
    return r.stdout.strip() if r.returncode == 0 else None


def get_settings():
    # brightness via GNOME session bus (no root)
    bright = None
    r = _run(["gdbus", "call", "--session", "--dest",
                        "org.gnome.SettingsDaemon.Power", "--object-path",
                        "/org/gnome/SettingsDaemon/Power", "--method",
                        "org.freedesktop.DBus.Properties.Get",
                        "org.gnome.SettingsDaemon.Power.Screen", "Brightness"],
                       capture_output=True, text=True)
    m = re.search(r"<\s*(?:int32\s+)?(-?\d+)", r.stdout)
    if m and int(m.group(1)) >= 0:
        bright = int(m.group(1))
    vol, muted = None, None
    rv = _run(["pactl", "get-sink-volume", "@DEFAULT_SINK@"],
                        capture_output=True, text=True)
    mv = re.search(r"(\d+)%", rv.stdout)
    if mv:
        vol = int(mv.group(1))
    rm = _run(["pactl", "get-sink-mute", "@DEFAULT_SINK@"],
                        capture_output=True, text=True)
    muted = "yes" in rm.stdout
    bt = None
    rb = _run(["bluetoothctl", "show"], capture_output=True,
                        text=True, timeout=6)
    if "Powered: yes" in rb.stdout:
        bt = True
    elif "Powered: no" in rb.stdout:
        bt = False
    wifi = None
    rw = _run(["nmcli", "radio", "wifi"], capture_output=True,
                        text=True)
    if rw.returncode == 0:
        wifi = rw.stdout.strip() == "enabled"
    return {
        "brightness": bright,
        "volume": vol, "muted": muted,
        "bluetooth": bt, "wifi": wifi,
        "theme": _gset(_IFACE, "color-scheme"),
        "wallpaper": _gset("org.gnome.desktop.background", "picture-uri"),
        "power_profile": _power_profile(),
        "power_profiles": ["power-saver", "balanced", "performance"],
        "night_light": _gbool(_COLOR, "night-light-enabled"),
        "night_temp": _gnum(_COLOR, "night-light-temperature"),
        "dnd": (lambda b: None if b is None else not b)(
            _gbool("org.gnome.desktop.notifications", "show-banners")),
        "battery_pct": _gbool(_IFACE, "show-battery-percentage"),
        "tap_click": _gbool(_TOUCHPAD, "tap-to-click"),
        "natural_scroll": _gbool(_TOUCHPAD, "natural-scroll"),
        "text_scale": _gnum(_IFACE, "text-scaling-factor"),
        "idle_blank": _gnum("org.gnome.desktop.session", "idle-delay"),
        "suspend_ac": _gnum(_POWER, "sleep-inactive-ac-timeout"),
        "slideshow": slideshow_state(),
        # Tweaks (GNOME Tweaks-style)
        "gtk_theme": _gset(_IFACE, "gtk-theme"),
        "icon_theme": _gset(_IFACE, "icon-theme"),
        "cursor_theme": _gset(_IFACE, "cursor-theme"),
        "font_name": _gset(_IFACE, "font-name"),
        "mono_font": _gset(_IFACE, "monospace-font-name"),
        "doc_font": _gset(_IFACE, "document-font-name"),
        "font_aa": _gset(_IFACE, "font-antialiasing"),
        "font_hint": _gset(_IFACE, "font-hinting"),
        "animations": _gbool(_IFACE, "enable-animations"),
        "hot_corner": _gbool(_IFACE, "enable-hot-corners"),
        "clock_weekday": _gbool(_IFACE, "clock-show-weekday"),
        "clock_date": _gbool(_IFACE, "clock-show-date"),
        "clock_seconds": _gbool(_IFACE, "clock-show-seconds"),
        "titlebar_buttons": _gset(_WM, "button-layout"),
        "ws_dynamic": _gbool(_MUTTER, "dynamic-workspaces"),
        "ws_num": _gnum(_WM, "num-workspaces"),
        "mouse_speed": _gnum(_MOUSE, "speed"),
        "touchpad_speed": _gnum(_TOUCHPAD, "speed"),
        "tweaks": _tweak_options(),
    }


def set_setting(key, value):
    if key == "brightness":
        v = max(5, min(100, int(value)))
        _run(["gdbus", "call", "--session", "--dest",
                        "org.gnome.SettingsDaemon.Power", "--object-path",
                        "/org/gnome/SettingsDaemon/Power", "--method",
                        "org.freedesktop.DBus.Properties.Set",
                        "org.gnome.SettingsDaemon.Power.Screen", "Brightness",
                        f"<int32 {v}>"], capture_output=True, timeout=6)
        return {"brightness": v}
    if key == "volume":
        v = max(0, min(150, int(value)))
        _run(["pactl", "set-sink-volume", "@DEFAULT_SINK@",
                        f"{v}%"], capture_output=True)
        return {"volume": v}
    if key == "mute":
        _run(["pactl", "set-sink-mute", "@DEFAULT_SINK@",
                        "toggle"], capture_output=True)
        return {"ok": True}
    if key == "theme":
        if value not in ("prefer-dark", "prefer-light", "default"):
            raise ValueError("bad theme")
        _run(["gsettings", "set", "org.gnome.desktop.interface",
                        "color-scheme", value], capture_output=True)
        _run(["gsettings", "set", "org.gnome.desktop.interface",
                        "gtk-theme",
                        "Yaru-dark" if value == "prefer-dark" else "Yaru"],
                       capture_output=True)
        return {"theme": value}
    if key == "bluetooth":
        _run(["bluetoothctl", "power",
                        "on" if value else "off"], capture_output=True,
                       timeout=6)
        return {"bluetooth": bool(value)}
    if key == "wifi":
        _run(["nmcli", "radio", "wifi",
                        "on" if value else "off"], capture_output=True)
        return {"wifi": bool(value)}
    if key == "wallpaper":
        path = os.path.realpath(value)
        if not os.path.isfile(path):
            raise ValueError("no such image")
        if preview_kind(path) != "image":
            raise ValueError("not an image file")
        uri = "file://" + urllib.parse.quote(path)
        for k in ("picture-uri", "picture-uri-dark"):
            _run(["gsettings", "set",
                            "org.gnome.desktop.background", k, uri],
                           capture_output=True)
        return {"wallpaper": uri}
    if key == "power_profile":
        if value not in ("power-saver", "balanced", "performance"):
            raise ValueError("bad power profile")
        _run(["powerprofilesctl", "set", value], capture_output=True)
        return {"power_profile": value}
    if key == "night_light":
        _gset_write(_COLOR, "night-light-enabled",
                    "true" if value else "false")
        return {"night_light": bool(value)}
    if key == "night_temp":
        _gset_write(_COLOR, "night-light-temperature",
                    f"uint32 {max(1700, min(6500, int(value)))}")
        return {"night_temp": int(value)}
    if key == "dnd":
        _gset_write("org.gnome.desktop.notifications", "show-banners",
                    "false" if value else "true")
        return {"dnd": bool(value)}
    if key == "battery_pct":
        _gset_write(_IFACE, "show-battery-percentage",
                    "true" if value else "false")
        return {"battery_pct": bool(value)}
    if key == "tap_click":
        _gset_write(_TOUCHPAD, "tap-to-click", "true" if value else "false")
        return {"tap_click": bool(value)}
    if key == "natural_scroll":
        _gset_write(_TOUCHPAD, "natural-scroll", "true" if value else "false")
        return {"natural_scroll": bool(value)}
    if key == "text_scale":
        _gset_write(_IFACE, "text-scaling-factor",
                    str(max(0.5, min(2.0, float(value)))))
        return {"text_scale": float(value)}
    if key == "idle_blank":
        _gset_write("org.gnome.desktop.session", "idle-delay",
                    f"uint32 {max(0, int(value))}")
        return {"idle_blank": int(value)}
    if key == "suspend_ac":
        secs = max(0, int(value))
        _gset_write(_POWER, "sleep-inactive-ac-timeout", str(secs))
        _gset_write(_POWER, "sleep-inactive-ac-type",
                    "nothing" if secs == 0 else "suspend")
        return {"suspend_ac": secs}
    if key in _TWEAK_BOOLS:
        sch, k = _TWEAK_BOOLS[key]
        _gset_write(sch, k, "true" if value else "false")
        return {key: bool(value)}
    if key in _TWEAK_CHOICES:
        sch, k, allowed = _TWEAK_CHOICES[key]
        if value not in allowed:
            raise ValueError("bad " + key)
        _gset_write(sch, k, value)
        return {key: value}
    if key in _TWEAK_THEMES:
        sch, k, opt = _TWEAK_THEMES[key]
        if value not in _tweak_options()[opt]:
            raise ValueError("theme not installed")
        _gset_write(sch, k, value)
        return {key: value}
    if key in _TWEAK_FONTS:
        v = str(value).strip()
        if not v or len(v) > 80:
            raise ValueError("bad font")
        _gset_write(_IFACE, _TWEAK_FONTS[key], v)
        return {key: v}
    if key == "titlebar_buttons":
        if value not in _TITLEBAR_LAYOUTS:
            raise ValueError("bad button layout")
        _gset_write(_WM, "button-layout", _TITLEBAR_LAYOUTS[value])
        return {key: value}
    if key == "ws_num":
        n = max(1, min(10, int(value)))
        _gset_write(_WM, "num-workspaces", str(n))
        return {key: n}
    if key in ("mouse_speed", "touchpad_speed"):
        v = max(-1.0, min(1.0, float(value)))
        _gset_write(_MOUSE if key == "mouse_speed" else _TOUCHPAD,
                    "speed", str(v))
        return {key: v}
    raise ValueError("unknown setting")


# ----------------------------------------------------- wallpaper slideshow ---

SLIDESHOW_FILE = os.path.join(CFG_DIR, "slideshow.json")
_slide_stop = threading.Event()
_slide_thread = [None]


def _slide_cfg():
    try:
        with open(SLIDESHOW_FILE) as f:
            c = json.load(f)
    except (OSError, ValueError):
        c = {}
    return {"enabled": bool(c.get("enabled")),
            "folder": c.get("folder", os.path.join(HOME, "Pictures")),
            "interval": int(c.get("interval", 300)),
            "shuffle": bool(c.get("shuffle", True))}


def _slide_images(folder):
    try:
        return sorted(os.path.join(folder, f) for f in os.listdir(folder)
                      if preview_kind(f) == "image")
    except OSError:
        return []


def _slide_loop():
    idx = 0
    while not _slide_stop.wait(1):
        cfg = _slide_cfg()
        if not cfg["enabled"]:
            return
        imgs = _slide_images(cfg["folder"])
        if not imgs:
            _slide_stop.wait(cfg["interval"])
            continue
        if cfg["shuffle"]:
            pick = imgs[secrets.randbelow(len(imgs))]
        else:
            pick = imgs[idx % len(imgs)]
            idx += 1
        uri = "file://" + urllib.parse.quote(pick)
        for k in ("picture-uri", "picture-uri-dark"):
            _gset_write("org.gnome.desktop.background", k, uri)
        _slide_stop.wait(cfg["interval"])


def _slide_start():
    _slide_stop.clear()
    if not (_slide_thread[0] and _slide_thread[0].is_alive()):
        _slide_thread[0] = threading.Thread(target=_slide_loop, daemon=True)
        _slide_thread[0].start()


def slideshow_state():
    c = _slide_cfg()
    c["running"] = bool(_slide_thread[0] and _slide_thread[0].is_alive())
    c["count"] = len(_slide_images(c["folder"]))
    return c


# ------------------------------------------------- backup & maintenance ------
# Simple rsync folder backups plus a weekly tidy-up job, both driven by
# systemd *user* timers written to ~/.config/systemd/user.

BACKUP_CFG = os.path.join(CFG_DIR, "backup.json")
BACKUP_STATE = os.path.join(MON_DIR, "backup-state.json")


def _user_unit_dir():
    d = os.path.join(HOME, ".config/systemd/user")
    os.makedirs(d, exist_ok=True)
    return d


def _install_user_timer(name, desc, exec_cmd, calendar):
    d = _user_unit_dir()
    atomic_write(os.path.join(d, name + ".service"),
                 f"[Unit]\nDescription={desc}\n\n[Service]\nType=oneshot\n"
                 f"ExecStart={exec_cmd}\n", 0o644)
    atomic_write(os.path.join(d, name + ".timer"),
                 f"[Unit]\nDescription={desc} (timer)\n\n[Timer]\n"
                 f"OnCalendar={calendar}\nPersistent=true\n\n"
                 "[Install]\nWantedBy=timers.target\n", 0o644)
    _run(["systemctl", "--user", "daemon-reload"])
    _run(["systemctl", "--user", "enable", "--now", name + ".timer"])


def _remove_user_timer(name):
    _run(["systemctl", "--user", "disable", "--now", name + ".timer"])
    d = _user_unit_dir()
    for ext in (".service", ".timer"):
        try:
            os.remove(os.path.join(d, name + ext))
        except OSError:
            pass
    _run(["systemctl", "--user", "daemon-reload"])


def _timer_state(name):
    en = _run(["systemctl", "--user", "is-enabled", name + ".timer"],
              capture_output=True, text=True)
    enabled = en.stdout.strip() == "enabled"
    nxt = ""
    if enabled:
        r = _run(["systemctl", "--user", "list-timers", name + ".timer",
                  "--no-legend"], capture_output=True, text=True)
        parts = r.stdout.split()
        nxt = " ".join(parts[:4]) if parts else ""
    return {"enabled": enabled, "next": nxt}


def backup_cfg():
    cfg = {"sources": [], "dest": "", "schedule": "off"}
    try:
        with open(BACKUP_CFG) as f:
            cfg.update(json.load(f))
    except (OSError, ValueError):
        pass
    return cfg


def _backup_inner(cfg):
    import shlex
    srcs = " ".join(shlex.quote(s) for s in cfg["sources"])
    dest = shlex.quote(cfg["dest"])
    state = shlex.quote(BACKUP_STATE)
    return (f"rsync -a --relative {srcs} {dest} && "
            "printf '{\"t\": %s, \"ok\": true}' \"$(date +%s)\" > " + state)


def backup_get():
    st = {}
    try:
        with open(BACKUP_STATE) as f:
            st = json.load(f)
    except (OSError, ValueError):
        pass
    return {**backup_cfg(), "last": st,
            "timer": _timer_state("perch-backup"),
            "rsync": bool(shutil.which("rsync"))}


def backup_set(body):
    sources = [os.path.realpath(os.path.expanduser(p.strip()))
               for p in body.get("sources", []) if p.strip()][:20]
    missing = [p for p in sources if not os.path.exists(p)]
    if missing:
        raise ValueError("source not found: " + missing[0])
    dest = os.path.realpath(os.path.expanduser(body.get("dest", "").strip()))
    schedule = body.get("schedule", "off")
    if schedule not in ("off", "daily", "weekly"):
        raise ValueError("bad schedule")
    if schedule != "off":
        if not sources:
            raise ValueError("add at least one source folder")
        if not os.path.isdir(dest):
            raise ValueError("destination folder does not exist")
        if any(dest == s or dest.startswith(s + os.sep) for s in sources):
            raise ValueError("destination must be outside the source folders")
    cfg = {"sources": sources, "dest": dest, "schedule": schedule}
    os.makedirs(CFG_DIR, exist_ok=True)
    atomic_write(BACKUP_CFG, json.dumps(cfg))
    if schedule == "off":
        _remove_user_timer("perch-backup")
    else:
        _install_user_timer("perch-backup", "Perch backup",
                            "/bin/sh -c " + _sh_quote(_backup_inner(cfg)),
                            schedule)
    return backup_get()


def _sh_quote(s):
    import shlex
    return shlex.quote(s)


def backup_run():
    cfg = backup_cfg()
    if not cfg["sources"]:
        raise ValueError("add at least one source folder")
    if not os.path.isdir(cfg["dest"]):
        raise ValueError("destination folder does not exist")
    if not shutil.which("rsync"):
        raise ValueError("rsync is not installed")
    return start_job(["sh", "-c", _backup_inner(cfg)], "Backup now")


_MAINT_CMD = ("/bin/sh -c 'rm -rf \"$HOME/.cache/thumbnails\"/* ; "
              "python3 -m pip cache purge >/dev/null 2>&1 ; "
              "notify-send \"Perch\" \"Weekly tidy-up done: thumbnail + pip "
              "caches cleared\" 2>/dev/null ; true'")


def maintenance_get():
    return _timer_state("perch-maintenance")


def maintenance_set(enabled):
    if enabled:
        _install_user_timer("perch-maintenance", "Perch weekly tidy-up",
                            _MAINT_CMD, "weekly")
    else:
        _remove_user_timer("perch-maintenance")
    return maintenance_get()


# ----------------------------------------------------------- db browser ------
# SQLite via the stdlib; Postgres via psql (host or `docker exec` into a
# running container). Read-only by default: only SELECT/PRAGMA/EXPLAIN/WITH
# statements run unless the request explicitly allows writes.

_READ_ONLY_SQL = re.compile(r"^\s*(select|pragma|explain|with|show|\\d)",
                            re.IGNORECASE)


def _guard_sql(sql, allow_write):
    if not sql.strip():
        raise ValueError("empty query")
    if not allow_write and not _READ_ONLY_SQL.match(sql):
        raise ValueError("read-only: only SELECT/PRAGMA/EXPLAIN/WITH allowed "
                         "(enable writes to run this)")


def _sqlite_tables(conn):
    cur = conn.execute("SELECT name FROM sqlite_master WHERE type IN "
                       "('table','view') ORDER BY name")
    return [r[0] for r in cur.fetchall()]


def sqlite_query(path, sql, allow_write=False, limit=500):
    import sqlite3
    path = os.path.realpath(os.path.expanduser(path))
    if not os.path.isfile(path):
        raise ValueError("no such database file")
    _guard_sql(sql, allow_write)
    mode = "ro" if not allow_write else "rw"
    uri = "file:" + urllib.parse.quote(path) + f"?mode={mode}"
    conn = sqlite3.connect(uri, uri=True, timeout=5)
    try:
        cur = conn.execute(sql)
        cols = [d[0] for d in cur.description] if cur.description else []
        rows = [list(r) for r in cur.fetchmany(limit)] if cols else []
        if allow_write:
            conn.commit()
        return {"columns": cols, "rows": rows,
                "tables": _sqlite_tables(conn),
                "rowcount": cur.rowcount if not cols else len(rows)}
    finally:
        conn.close()


def _psql_argv(conn_str, container):
    if container:
        if not re.fullmatch(r"[a-zA-Z0-9][\w.-]{0,80}", container):
            raise ValueError("bad container name")
        return ["docker", "exec", "-i", container, "psql", conn_str]
    return ["psql", conn_str]


def pg_query(conn_str, sql, container="", allow_write=False, limit=500):
    if not shutil.which("docker" if container else "psql"):
        raise ValueError("psql not available"
                         + (" and no docker" if container else ""))
    _guard_sql(sql, allow_write)
    # cap read-only result sets that don't already carry their own LIMIT
    wrapped = f"SELECT * FROM (\n{sql}\n) AS q LIMIT {int(limit)}" \
        if _READ_ONLY_SQL.match(sql) and "limit" not in sql.lower() else sql
    # -A unaligned output, -F tab field separator, footer off
    r = _run(_psql_argv(conn_str, container)
             + ["-A", "-F", "\t", "-P", "footer=off", "-c", wrapped],
             capture_output=True, text=True, timeout=20,
             env={**os.environ, "PGCONNECT_TIMEOUT": "5"})
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or "query failed")[:300])
    lines = r.stdout.splitlines()
    if not lines:
        return {"columns": [], "rows": [], "rowcount": 0}
    cols = lines[0].split("\t")
    rows = [ln.split("\t") for ln in lines[1:]]
    return {"columns": cols, "rows": rows, "rowcount": len(rows)}


def crontab_get():
    r = _run(["crontab", "-l"], capture_output=True, text=True)
    text = r.stdout if r.returncode == 0 else ""
    entries = []
    for line in text.splitlines():
        s = line.strip()
        if s and not s.startswith("#") and "=" not in s.split()[0]:
            parts = s.split(None, 5)
            if len(parts) >= 6:
                entries.append({"schedule": " ".join(parts[:5]),
                                "command": parts[5]})
    return {"raw": text, "entries": entries,
            "installed": bool(shutil.which("crontab"))}


def crontab_set(text):
    if len(text) > 64000:
        raise ValueError("crontab too large")
    p = subprocess.Popen(["crontab", "-"], stdin=subprocess.PIPE,
                         stderr=subprocess.PIPE, text=True)
    _, err = p.communicate(text if text.endswith("\n") else text + "\n",
                          timeout=10)
    if p.returncode != 0:
        raise ValueError((err.strip() or "crontab rejected")[:300])
    return crontab_get()


def timers_list():
    out = []
    for scope in ("--user", "--system"):
        r = _run(["systemctl", scope, "list-timers", "--all", "--no-pager",
                  "--output=json"], capture_output=True, text=True)
        if r.returncode != 0:
            continue
        try:
            data = json.loads(r.stdout or "[]")
        except ValueError:
            continue
        for t in data:
            unit = t.get("unit", "")
            en = _run(["systemctl", scope, "is-enabled", unit],
                      capture_output=True, text=True)
            out.append({"unit": unit, "scope": scope.strip("-"),
                        "next": t.get("next", "") or "",
                        "left": t.get("left", "") or "",
                        "activates": t.get("activates", ""),
                        "enabled": en.stdout.strip() == "enabled"})
    return {"timers": out}


_TIMER_RE = re.compile(r"^[\w@.\\:-]+\.timer$")


def timer_action(unit, action):
    if action not in ("enable", "disable", "start", "stop") \
            or not _TIMER_RE.fullmatch(unit):
        raise ValueError("bad timer action")
    args = ["--now"] if action in ("enable", "disable") else []
    r = _run(["systemctl", "--user", action, *args, unit],
             capture_output=True, text=True, timeout=20)
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or "failed")[:300])
    return timers_list()


SSH_DIR = os.path.join(HOME, ".ssh")


def ssh_keys():
    keys = []
    try:
        names = sorted(os.listdir(SSH_DIR))
    except OSError:
        names = []
    for name in names:
        pub = os.path.join(SSH_DIR, name)
        if not name.endswith(".pub") or not os.path.isfile(pub):
            continue
        with open(pub) as f:
            content = f.read().strip()
        r = _run(["ssh-keygen", "-lf", pub], capture_output=True, text=True)
        fp = r.stdout.strip() if r.returncode == 0 else ""
        keys.append({"name": name, "public": content, "fingerprint": fp,
                     "has_private": os.path.isfile(pub[:-4])})
    auth = ""
    try:
        with open(os.path.join(SSH_DIR, "authorized_keys")) as f:
            auth = f.read()
    except OSError:
        pass
    return {"keys": keys, "authorized_keys": auth,
            "installed": bool(shutil.which("ssh-keygen"))}


def ssh_keygen(name, comment):
    if not re.fullmatch(r"[a-zA-Z0-9._-]{1,40}", name):
        raise ValueError("bad key name")
    if not shutil.which("ssh-keygen"):
        raise ValueError("ssh-keygen is not installed")
    os.makedirs(SSH_DIR, exist_ok=True)
    os.chmod(SSH_DIR, 0o700)
    path = os.path.join(SSH_DIR, name)
    if os.path.exists(path) or os.path.exists(path + ".pub"):
        raise ValueError("a key with that name already exists")
    who = pwd.getpwuid(os.getuid()).pw_name
    comment = re.sub(r"[^\w@.\-]", "", comment)[:60] or f"{who}@perch"
    r = _run(["ssh-keygen", "-t", "ed25519", "-f", path, "-N", "",
              "-C", comment], capture_output=True, text=True)
    if r.returncode != 0:
        raise ValueError((r.stderr.strip() or "ssh-keygen failed")[:300])
    return ssh_keys()


def pg_containers():
    """Running containers whose image looks like Postgres — for a quick picker."""
    if not shutil.which("docker"):
        return []
    r = _run(["docker", "ps", "--format", "{{.Names}}\t{{.Image}}"],
             capture_output=True, text=True)
    out = []
    for line in r.stdout.splitlines():
        parts = line.split("\t")
        if len(parts) == 2 and ("postgres" in parts[1].lower()
                                or "postgis" in parts[1].lower()):
            out.append({"name": parts[0], "image": parts[1]})
    return out


def slideshow_set(cfg):
    os.makedirs(CFG_DIR, exist_ok=True)
    cur = _slide_cfg()
    if "folder" in cfg:
        cur["folder"] = os.path.realpath(os.path.expanduser(cfg["folder"]))
    if "interval" in cfg:
        cur["interval"] = max(5, int(cfg["interval"]))
    if "shuffle" in cfg:
        cur["shuffle"] = bool(cfg["shuffle"])
    cur["enabled"] = bool(cfg.get("enabled", cur["enabled"]))
    atomic_write(SLIDESHOW_FILE, json.dumps(cur))
    if cur["enabled"]:
        _slide_start()
    else:
        _slide_stop.set()
    return slideshow_state()


if _slide_cfg()["enabled"]:
    _slide_start()


# -------------------------------------------------------- open-with / IDE ----

APP_TABLE = [
    ("code", "VS Code", ["code"], "both"),
    ("subl", "Sublime Text", ["subl"], "both"),
    ("idea", "IntelliJ IDEA", ["idea"], "both"),
    ("pycharm", "PyCharm", ["pycharm"], "both"),
    ("goland", "GoLand", ["goland"], "both"),
    ("webstorm", "WebStorm", ["webstorm"], "both"),
    ("gnome-text-editor", "Text Editor", ["gnome-text-editor"], "file"),
    ("gedit", "gedit", ["gedit"], "file"),
    ("nautilus", "Files (Nautilus)", ["nautilus"], "dir"),
    ("xdg-open", "Default app", ["xdg-open"], "both"),
]


def available_apps():
    out = []
    for key, label, cmd, kind in APP_TABLE:
        if shutil.which(cmd[0]):
            out.append({"key": key, "label": label, "kind": kind})
    return out


def open_with(app, path):
    path = os.path.realpath(path)
    if not os.path.exists(path):
        raise ValueError("no such path")
    for key, label, cmd, kind in APP_TABLE:
        if key == app and shutil.which(cmd[0]):
            subprocess.Popen([*cmd, path], stdout=subprocess.DEVNULL,
                             stderr=subprocess.DEVNULL)
            return label
    raise ValueError("app not available")


# --------------------------------------------------------- site preview ------

SHOT_DIR = os.path.join(MON_DIR, "shots")


def site_shot(url):
    if not re.match(r"^https?://", url or ""):
        url = "https://" + (url or "").strip()
    if not re.match(r"^https?://[\w.-]", url):
        raise ValueError("invalid URL")
    os.makedirs(SHOT_DIR, exist_ok=True)
    name = re.sub(r"\W+", "_", url)[:60] + ".png"
    out = os.path.join(SHOT_DIR, name)
    chrome = shutil.which("google-chrome") or shutil.which("chromium-browser")
    if not chrome:
        raise ValueError("chrome not installed")
    r = _run([chrome, "--headless=new", "--disable-gpu",
                        "--hide-scrollbars", "--window-size=1280,1400",
                        f"--screenshot={out}", "--virtual-time-budget=6000",
                        url], timeout=90, capture_output=True, text=True)
    if not os.path.exists(out):
        raise ValueError((r.stderr.strip()[-200:]) or "screenshot failed")
    return {"path": out, "url": url}


# ---------------------------------------------------------- runtimes ---------

RUNTIMES = [
    ("Node.js", ["node", "-v"]), ("npm", ["npm", "-v"]),
    ("Python", ["python3", "--version"]),
    ("Java", ["java", "-version"]), ("Go", ["go", "version"]),
    ("Rust", ["rustc", "--version"]), ("Ruby", ["ruby", "-v"]),
    ("PHP", ["php", "-v"]), ("Deno", ["deno", "-V"]),
]


def runtimes():
    out = []
    for label, cmd in RUNTIMES:
        if not shutil.which(cmd[0]):
            continue
        try:
            r = _run(cmd, capture_output=True, text=True, timeout=8)
            line = (r.stdout or r.stderr).strip().splitlines()
            out.append({"name": label, "version": line[0][:70] if line else "?",
                        "path": shutil.which(cmd[0])})
        except Exception:  # noqa: BLE001
            continue
    rust_tc = []
    rust_active = None
    if shutil.which("rustup"):
        r = _run(["rustup", "toolchain", "list"],
                           capture_output=True, text=True)
        for ln in r.stdout.splitlines():
            tc = ln.split()[0]
            rust_tc.append(tc)
            if "default" in ln:
                rust_active = tc
    node_versions = []
    nvm = os.path.join(HOME, ".nvm/versions/node")
    if os.path.isdir(nvm):
        node_versions = sorted(os.listdir(nvm), reverse=True)
    return {"tools": out, "rust_toolchains": rust_tc,
            "rust_active": rust_active, "node_versions": node_versions,
            "has_nvm": os.path.exists(os.path.join(HOME, ".nvm/nvm.sh")),
            "alternatives": alternatives()}


def alternatives():
    """Debian update-alternatives groups that have more than one candidate —
    the general 'switch the active version of a tool' mechanism (java, python,
    editor, browsers, …). Switching needs root, so it runs via pkexec."""
    groups = []
    r = _run(["update-alternatives", "--get-selections"],
                       capture_output=True, text=True)
    for line in r.stdout.splitlines():
        parts = line.split()
        if len(parts) < 3:
            continue
        name, mode, current = parts[0], parts[1], parts[2]
        rl = _run(["update-alternatives", "--list", name],
                            capture_output=True, text=True)
        opts = [o for o in rl.stdout.splitlines() if o]
        if len(opts) > 1:
            groups.append({"name": name, "current": current,
                           "auto": mode == "auto", "options": opts})
    groups.sort(key=lambda g: g["name"])
    return groups


def set_runtime(kind, value):
    if kind == "rust":
        if not re.fullmatch(r"[\w.+-]+", value):
            raise ValueError("bad toolchain")
        return start_job(["rustup", "default", value],
                         f"Set Rust default → {value}")
    raise ValueError("switching this runtime isn't supported here")


# --------------------------------------------------------- office docs -------


def read_office(path):
    path = os.path.realpath(path)
    ext = os.path.splitext(path.lower())[1]
    if ext == ".docx":
        import docx
        d = docx.Document(path)
        blocks = []
        for para in d.paragraphs:
            if para.text.strip():
                style = (para.style.name or "").lower()
                blocks.append({"h": "heading" in style, "text": para.text})
        tables = [[[c.text for c in row.cells] for row in t.rows[:60]]
                  for t in d.tables[:8]]
        return {"kind": "docx", "blocks": blocks[:800], "tables": tables}
    if ext in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheets = []
        for ws in wb.worksheets[:12]:
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 200:
                    break
                rows.append(["" if c is None else str(c)[:200]
                             for c in row[:40]])
            try:
                dims = ws.calculate_dimension()
            except Exception:  # noqa: BLE001
                dims = f"{len(rows)} rows"
            sheets.append({"name": ws.title, "rows": rows, "dims": dims})
        wb.close()
        return {"kind": "xlsx", "sheets": sheets}
    raise ValueError("unsupported office format")


# ------------------------------------------------------------ file editor ----

MAX_EDIT = 2 * 2 ** 20  # 2 MB


def read_text_file(path):
    path = os.path.realpath(path)
    if not os.path.isfile(path):
        raise ValueError("not a file")
    if os.path.getsize(path) > MAX_EDIT:
        raise ValueError("file larger than 2 MB — open it in a real editor")
    with open(path, "rb") as f:
        raw = f.read()
    if b"\x00" in raw[:8192]:
        raise ValueError("binary file — use Open instead")
    return {"path": path, "content": raw.decode("utf-8", errors="replace"),
            "writable": os.access(path, os.W_OK)}


def _check_home_write(path):
    path = os.path.realpath(path)
    if not path.startswith(HOME + os.sep):
        raise ValueError("writing is limited to your home directory")
    if not os.path.isdir(os.path.dirname(path)):
        raise ValueError("folder does not exist")
    return path


def write_text_file(path, content):
    path = _check_home_write(path)
    if len(content) > MAX_EDIT:
        raise ValueError("content too large (2 MB limit)")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def save_png(path, data_b64):
    import base64
    path = _check_home_write(path)
    if not path.endswith(".png"):
        path += ".png"
    raw = base64.b64decode(data_b64.split(",")[-1])
    if len(raw) > 8 * 2 ** 20 or not raw.startswith(b"\x89PNG"):
        raise ValueError("not a valid PNG")
    with open(path, "wb") as f:
        f.write(raw)
    return path


# ------------------------------------------------------------------ server ---


# ------------------------------------------------------ websocket terminal ---
# Minimal RFC 6455 server (no fragmentation/extensions — plenty for a
# terminal) feeding a pty running the user's shell. One pty per connection.

_WS_GUID = "258EAFA5-E914-47DA-95CA-C5AB0DC85B11"


def _ws_accept(key):
    return base64.b64encode(
        hashlib.sha1((key + _WS_GUID).encode()).digest()).decode()


def _ws_send(sock, data, opcode=0x1):
    payload = data.encode() if isinstance(data, str) else data
    head = bytes([0x80 | opcode])
    n = len(payload)
    if n < 126:
        head += bytes([n])
    elif n < 65536:
        head += bytes([126]) + struct.pack(">H", n)
    else:
        head += bytes([127]) + struct.pack(">Q", n)
    sock.sendall(head + payload)


def _ws_recv(sock):
    """Read one client frame → (opcode, payload) or (None, b'') on EOF."""
    def read(n):
        buf = b""
        while len(buf) < n:
            chunk = sock.recv(n - len(buf))
            if not chunk:
                raise ConnectionError
            buf += chunk
        return buf
    try:
        b1, b2 = read(2)
    except (ConnectionError, OSError):
        return None, b""
    opcode = b1 & 0x0F
    masked = b2 & 0x80
    n = b2 & 0x7F
    if n == 126:
        n = struct.unpack(">H", read(2))[0]
    elif n == 127:
        n = struct.unpack(">Q", read(8))[0]
    if n > 2 ** 20:
        return None, b""
    mask = read(4) if masked else b"\0\0\0\0"
    data = bytes(c ^ mask[i % 4] for i, c in enumerate(read(n)))
    return opcode, data


def _pty_session(sock, cwd=None, cmd=None):
    """Bridge a websocket to a fresh login shell in a pty until either side
    closes. Client protocol: 'i<data>' keyboard input, 'r{cols,rows}' resize.
    Optional cwd starts the shell in that directory; optional cmd is typed as
    the first line (used to `cd` here / `docker exec` into a container)."""
    shell = os.environ.get("SHELL") or pwd.getpwuid(os.getuid()).pw_shell \
        or "/bin/bash"
    start = HOME
    if cwd:
        start = cwd if os.path.isdir(cwd) else (
            os.path.dirname(cwd) if os.path.isfile(cwd) else HOME)
    pid, fd = _pty.fork()
    if pid == 0:  # child
        try:
            os.chdir(start)
        except OSError:
            os.chdir(HOME)
        env = {**os.environ, "TERM": "xterm-256color"}
        os.execve(shell, [shell, "-l"], env)
    if cmd:  # run an initial command in the freshly-started shell
        os.write(fd, (cmd + "\n").encode())
    try:
        while True:
            r, _, _ = select.select([fd, sock], [], [], 30)
            if fd in r:
                try:
                    out = os.read(fd, 32768)
                except OSError:
                    break
                if not out:
                    break
                _ws_send(sock, out, opcode=0x2)
            if sock in r:
                op, data = _ws_recv(sock)
                if op is None or op == 0x8:
                    break
                if op == 0x9:  # ping → pong
                    _ws_send(sock, data, opcode=0xA)
                    continue
                if not data:
                    continue
                kind, rest = data[:1], data[1:]
                if kind == b"i":
                    os.write(fd, rest)
                elif kind == b"r":
                    try:
                        d = json.loads(rest)
                        winsz = struct.pack("HHHH", int(d["rows"]),
                                            int(d["cols"]), 0, 0)
                        fcntl.ioctl(fd, termios.TIOCSWINSZ, winsz)
                    except (ValueError, KeyError, OSError):
                        pass
    finally:
        try:
            os.close(fd)
        except OSError:
            pass
        try:
            os.kill(pid, 15)
            os.waitpid(pid, os.WNOHANG)
        except OSError:
            pass


class Handler(BaseHTTPRequestHandler):
    protocol_version = "HTTP/1.1"

    def log_message(self, *a):
        pass

    def _send(self, code, body, ctype="application/json"):
        data = body if isinstance(body, bytes) else body.encode()
        self.send_response(code)
        self.send_header("Content-Type", ctype + "; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(data)

    def _json(self, obj, code=200):
        self._send(code, json.dumps(obj))

    def _err(self, msg, code=400):
        self._json({"error": str(msg)}, code)

    def _raw(self, path):
        path = os.path.realpath(path)
        if not os.path.isfile(path):
            return self._err("no such file", 404)
        size = os.path.getsize(path)
        if size > 512 * 2 ** 20:
            return self._err("file larger than 512 MB", 413)
        ext = os.path.splitext(path.lower())[1]
        ctype = MIMES.get(ext, "application/octet-stream")
        try:
            f = open(path, "rb")
        except OSError as e:
            return self._err(e, 403)
        with f:
            self.send_response(200)
            self.send_header("Content-Type", ctype)
            self.send_header("Content-Length", str(size))
            self.send_header("Cache-Control", "no-store")
            self.end_headers()
            while chunk := f.read(256 * 1024):
                self.wfile.write(chunk)

    def _static(self, rel):
        # containment by realpath, not by stripping "..": a blocklist has to be
        # re-proved correct every time the input encoding changes
        root = os.path.realpath(os.path.join(WEB_DIR, "static"))
        path = os.path.realpath(os.path.join(root, rel.lstrip("/")))
        if path != root and not path.startswith(root + os.sep):
            return self._err("not found", 404)
        if not os.path.isfile(path):
            return self._err("not found", 404)
        with open(path, "rb") as f:
            data = f.read()
        ext = os.path.splitext(path)[1]
        self.send_response(200)
        self.send_header("Content-Type",
                         STATIC_TYPES.get(ext, "application/octet-stream"))
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-cache")
        self.end_headers()
        self.wfile.write(data)

    def _cookie_token(self):
        for part in self.headers.get("Cookie", "").split(";"):
            k, _, v = part.strip().partition("=")
            if k == "perch_t":
                return v
        return ""

    def _authed(self, qs):
        ip = self.client_address[0]
        if _auth_locked(ip):
            return False
        tok = (self.headers.get("X-Token") or qs.get("t", [""])[0]
               or self._cookie_token())
        if not secrets.compare_digest(tok, TOKEN):
            _auth_fail(ip)
            return False
        return True

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        qs = urllib.parse.parse_qs(parsed.query)
        route = parsed.path
        if route == "/api/health":
            return self._json({"ok": True})
        if route.startswith("/static/"):
            return self._static(route[len("/static/"):])
        if route == "/":
            if not self._authed(qs):
                self._send(403, "<h3>Forbidden — open the exact URL printed "
                                "at startup (it contains the access token).",
                           "text/html")
                return
            if qs.get("t"):
                # move the URL token into an HttpOnly cookie + clean the URL
                self.send_response(303)
                self.send_header("Set-Cookie", f"perch_t={TOKEN}; HttpOnly; "
                                 "SameSite=Strict; Path=/")
                self.send_header("Location", "/")
                self.send_header("Content-Length", "0")
                self.end_headers()
                return
            self._send(200, render_index(TOKEN), "text/html")
            return
        if not self._authed(qs):
            return self._err("bad token", 403)
        if route == "/ws/term":
            key = self.headers.get("Sec-WebSocket-Key")
            if (self.headers.get("Upgrade", "").lower() != "websocket"
                    or not key):
                return self._err("websocket endpoint", 400)
            self.send_response(101, "Switching Protocols")
            self.send_header("Upgrade", "websocket")
            self.send_header("Connection", "Upgrade")
            self.send_header("Sec-WebSocket-Accept", _ws_accept(key))
            self.end_headers()
            self.close_connection = True
            cwd = qs.get("cwd", [""])[0] or None
            if cwd:
                cwd = os.path.realpath(cwd)
                if not (cwd == HOME or cwd.startswith(HOME + os.sep)
                        or os.path.isdir(cwd)):
                    cwd = None
            cmd = None
            if qs.get("cmd", [""])[0]:
                try:
                    cmd = base64.b64decode(qs["cmd"][0]).decode()[:2000]
                except (ValueError, UnicodeDecodeError):
                    cmd = None
            try:
                _pty_session(self.connection, cwd=cwd, cmd=cmd)
            except (OSError, ConnectionError):
                pass
            return
        try:
            if route == "/api/overview":
                return self._json(overview())
            if route == "/api/history":
                return self._json(list(HISTORY))
            if route == "/api/disks":
                return self._json(disks())
            if route == "/api/processes":
                return self._json(processes(qs.get("sort", ["cpu"])[0],
                                            qs.get("q", [""])[0]))
            if route == "/api/users":
                return self._json(users())
            if route == "/api/browse":
                return self._json(browse(qs.get("path", [HOME])[0]))
            if route == "/api/du":
                return self._json(analyze(qs.get("path", [HOME])[0]))
            if route == "/api/cleanup":
                return self._json(cleanup_report())
            if route == "/api/search":
                return self._json(search_files(
                    qs.get("q", [""])[0],
                    regex=qs.get("regex", ["0"])[0] == "1"))
            if route == "/api/net":
                return self._json(net_ports())
            if route == "/api/docker":
                return self._json(docker_info())
            if route == "/api/dockerlogs":
                return self._json(docker_logs(qs.get("id", [""])[0]))
            if route == "/api/services":
                return self._json(services())
            if route == "/api/devinfo":
                return self._json(devinfo())
            if route == "/api/kernel":
                return self._json(kernel_info())
            if route == "/api/gpu":
                return self._json(gpu_info())
            if route == "/api/hw":
                return self._json(hw_info())
            if route == "/api/logs":
                return self._json(read_logs(qs.get("source", ["system"])[0],
                                            qs.get("n", ["150"])[0],
                                            qs.get("q", [""])[0],
                                            qs.get("prio", [""])[0],
                                            qs.get("cursor", [""])[0]))
            if route == "/api/readfile":
                return self._json(read_text_file(qs.get("path", [""])[0]))
            if route == "/api/monitor":
                return self._json(monitor_state(
                    qs.get("brief", ["0"])[0] == "1"))
            if route == "/api/notify":
                return self._json(notify_cfg())
            if route == "/api/logwatch":
                return self._json(logwatch_state())
            if route == "/api/caps":
                return self._json(capabilities())
            if route == "/api/healthscore":
                return self._json(health_score())
            if route == "/api/backup":
                return self._json(backup_get())
            if route == "/api/maintenance":
                return self._json(maintenance_get())
            if route == "/api/pgcontainers":
                return self._json({"containers": pg_containers()})
            if route == "/api/cron":
                return self._json(crontab_get())
            if route == "/api/timers":
                return self._json(timers_list())
            if route == "/api/sshkeys":
                return self._json(ssh_keys())
            if route == "/api/projectscripts":
                return self._json(project_scripts(qs.get("path", [""])[0]))
            if route == "/api/updates":
                return self._json(pkg_updates(qs.get("force", ["0"])[0] == "1"))
            if route == "/api/procinfo":
                return self._json(proc_detail(int(qs.get("pid", ["0"])[0])))
            if route == "/api/speedtest":
                return self._json(speed_test())
            if route == "/api/job":
                return self._json(job_status(qs.get("id", [""])[0],
                                             qs.get("since", ["0"])[0]))
            if route == "/api/pkgsearch":
                return self._json(pkg_search(qs.get("q", [""])[0]))
            if route == "/api/settings":
                return self._json(get_settings())
            if route == "/api/apps":
                return self._json({"apps": available_apps()})
            if route == "/api/runtimes":
                return self._json(runtimes())
            if route == "/api/llm":
                return self._json(llm_public())
            if route == "/api/office":
                return self._json(read_office(qs.get("path", [""])[0]))
            if route == "/api/shot":
                return self._raw(qs.get("path", [""])[0])
            if route == "/api/raw":
                return self._raw(qs.get("path", [""])[0])
            if route == "/api/httpstore":
                return self._json(http_store_get())
            if route == "/api/gitrepos":
                return self._json(git_repos())
            if route == "/api/dockerstats":
                return self._json(docker_stats())
            if route == "/api/dockercompose":
                return self._json(docker_compose_projects())
            if route == "/api/containers":
                return self._json(container_envs())
            if route == "/api/ctrlogs":
                return self._json(ctr_logs(qs.get("engine", [""])[0],
                                           qs.get("id", [""])[0],
                                           qs.get("ns", [""])[0]))
            if route == "/api/alertctl":
                return self._json(alert_ctl())
            if route == "/api/firewall":
                return self._json(firewall_status())
            if route == "/api/diskhealth":
                return self._json(disk_health())
            if route == "/api/dockerdisk":
                return self._json(docker_disk())
            if route == "/api/homelayout":
                return self._json(home_layout())
            if route == "/api/dupes":
                return self._json(find_duplicates(
                    qs.get("path", [""])[0], qs.get("min", ["1"])[0],
                    qs.get("secs", ["20"])[0]))
            if route == "/api/oldfiles":
                return self._json(find_old_large(
                    qs.get("path", [""])[0], qs.get("min", ["100"])[0],
                    qs.get("days", ["365"])[0], qs.get("secs", ["20"])[0]))
            return self._err("not found", 404)
        except Exception as e:  # noqa: BLE001 — surfaced to the UI
            return self._err(e)

    def do_POST(self):
        parsed = urllib.parse.urlparse(self.path)
        if not self._authed(urllib.parse.parse_qs(parsed.query)):
            return self._err("bad token", 403)
        try:
            n = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(n) or b"{}")
            route = parsed.path
            if route == "/api/kill":
                name = kill_proc(int(body["pid"]), bool(body.get("force")))
                return self._json({"ok": True, "name": name})
            if route == "/api/open":
                return self._json({"ok": True, "path": open_file(body["path"])})
            if route == "/api/trash":
                paths = body.get("paths") or [body["path"]]
                done, errors = [], []
                for p in paths[:500]:
                    try:
                        done.append(trash_file(p))
                    except Exception as e:  # noqa: BLE001
                        errors.append(f"{os.path.basename(p)}: {e}")
                return self._json({"ok": not errors, "trashed": len(done),
                                   "errors": errors[:5]})
            if route == "/api/clean":
                freed = do_clean(body["target"], body.get("path"))
                return self._json({"ok": True, "freed": freed})
            if route == "/api/reindex":
                threading.Thread(target=build_index, daemon=True).start()
                return self._json({"ok": True})
            if route == "/api/killport":
                pid = body.get("pid")
                return self._json({"ok": True,
                                   "name": kill_port(int(body["port"]),
                                                     bool(body.get("force")),
                                                     int(pid) if pid else None)})
            if route == "/api/dockeraction":
                out = docker_action(body["id"], body["action"])
                return self._json({"ok": True, "out": out})
            if route == "/api/serviceaction":
                service_action(body["name"], body["action"])
                return self._json({"ok": True})
            if route == "/api/editor":
                return self._json({"ok": True, "editor": open_editor(body["path"])})
            if route == "/api/writefile":
                p = write_text_file(body["path"], body.get("content", ""))
                return self._json({"ok": True, "path": p})
            if route == "/api/savepng":
                p = save_png(body["path"], body["data"])
                return self._json({"ok": True, "path": p})
            if route == "/api/ai":
                cfg = llm_cfg()
                prompt = body.get("prompt", "")
                if cfg["provider"] == "claude-cli":
                    return self._json(ask_ai(prompt,
                                             bool(body.get("snapshot", True)),
                                             bool(body.get("reset")),
                                             body.get("extra", "")))
                system = ("You are the assistant inside a local Linux system "
                          "dashboard. Current system snapshot:\n" + ai_snapshot()
                          ) if body.get("snapshot", True) else ""
                history = body.get("history") or []
                msgs = [{"role": m["role"], "content": m["content"]}
                        for m in history if m.get("role") in ("user", "assistant")]
                msgs.append({"role": "user", "content": prompt})
                text = _provider_chat(cfg, system, msgs, 2048)
                return self._json({"text": text, "provider": cfg["provider"]})
            if route == "/api/llmconfig":
                return self._json(llm_save(body))
            if route == "/api/llmtest":
                return self._json(llm_test())
            if route == "/api/setalternative":
                name, path = body["name"], body["path"]
                if not re.fullmatch(r"[\w.@+-]+", name) or not os.path.exists(path):
                    return self._err("bad alternative")
                return self._json(start_job(
                    ["update-alternatives", "--set", name, path],
                    f"Switch {name} → {os.path.basename(path)}",
                    privileged=True))
            if route == "/api/monitor":
                return self._json({"ok": True, "cfg": mon_save(body)})
            if route == "/api/testnotify":
                return self._json(notify_test())
            if route == "/api/notifyconfig":
                return self._json(notify_save(body))
            if route == "/api/logwatchconfig":
                return self._json(logwatch_save(body))
            if route == "/api/http":
                res = http_request(body.get("method"), body.get("url"),
                                   body.get("headers", ""), body.get("body", ""),
                                   body.get("timeout", 20))
                try:
                    http_history_add({"t": time.time(),
                                      "method": (body.get("method") or "GET"),
                                      "url": body.get("url"),
                                      "status": res["status"], "ms": res["ms"]})
                except Exception:  # noqa: BLE001
                    pass
                return self._json(res)
            if route == "/api/yaml":
                try:
                    return self._json({"ok": True, "text": yaml_convert(
                        body.get("text", ""), body.get("dir", "y2j"))})
                except Exception as e:  # noqa: BLE001
                    return self._err(f"conversion failed: {e}")
            if route == "/api/pkginstall":
                return self._json(pkg_install(body["mgr"], body["name"]))
            if route == "/api/upgradeall":
                return self._json(upgrade_all(body["mgr"]))
            if route == "/api/setsetting":
                return self._json(set_setting(body["key"], body.get("value")))
            if route == "/api/slideshow":
                return self._json(slideshow_set(body))
            if route == "/api/backup":
                return self._json(backup_set(body))
            if route == "/api/backuprun":
                return self._json(backup_run())
            if route == "/api/maintenance":
                return self._json(maintenance_set(bool(body.get("enabled"))))
            if route == "/api/dbquery":
                engine = body.get("engine")
                if engine == "sqlite":
                    return self._json(sqlite_query(
                        body.get("path", ""), body.get("sql", ""),
                        bool(body.get("write"))))
                if engine == "postgres":
                    return self._json(pg_query(
                        body.get("conn", ""), body.get("sql", ""),
                        body.get("container", ""), bool(body.get("write"))))
                return self._err("unknown engine")
            if route == "/api/cronsave":
                return self._json(crontab_set(body.get("text", "")))
            if route == "/api/timeraction":
                return self._json(timer_action(body["unit"], body["action"]))
            if route == "/api/sshkeygen":
                return self._json(ssh_keygen(body.get("name", ""),
                                             body.get("comment", "")))
            if route == "/api/openwith":
                return self._json({"ok": True,
                                   "app": open_with(body["app"], body["path"])})
            if route == "/api/sitepreview":
                return self._json(site_shot(body["url"]))
            if route == "/api/setruntime":
                return self._json(set_runtime(body["kind"], body["value"]))
            if route == "/api/httpstore":
                return self._json(http_store_save(body))
            if route == "/api/gitaction":
                return self._json(git_action(body["path"], body["action"]))
            if route == "/api/projectrun":
                return self._json(project_run(body["path"], body["kind"],
                                              body["name"]))
            if route == "/api/composeaction":
                return self._json(docker_compose_action(body["project"],
                                                        body["action"]))
            if route == "/api/dockerprune":
                return self._json(docker_prune(body["kind"]))
            if route == "/api/ctraction":
                out = ctr_action(body.get("engine", ""), body.get("id", ""),
                                 body.get("action", ""))
                return self._json({"ok": True, "out": out})
            if route == "/api/alertctl":
                return self._json({"ok": True,
                                   "ctl": alert_ctl_set(body.get("action", ""),
                                                        body.get("minutes", 0))})
            if route == "/api/customrules":
                return self._json({"ok": True,
                                   "rules": custom_save(body.get("rules", []))})
            if route == "/api/homelayout":
                return self._json(home_layout_save(body))
            if route == "/api/firewallrules":
                return self._json(firewall_rules_job())
            if route == "/api/smart":
                return self._json(smart_job(body.get("device", "")))
            if route == "/api/health":
                return self._json(health_report())
            return self._err("not found", 404)
        except psutil.AccessDenied:
            return self._err("permission denied — that process is not yours", 403)
        except psutil.NoSuchProcess:
            return self._err("process already gone", 410)
        except Exception as e:  # noqa: BLE001
            return self._err(e)


# -------------------------------------------------------------------- page ---




def _watchdog_loop():
    """Ping systemd's watchdog only while our own HTTP endpoint answers."""
    import urllib.request
    target = "127.0.0.1" if HOST in ("0.0.0.0", "::") else HOST
    while True:
        time.sleep(20)
        try:
            urllib.request.urlopen(
                f"http://{target}:{PORT}/api/health", timeout=5)
            _sd_notify("WATCHDOG=1")
        except OSError:
            pass  # skip the ping — systemd restarts us if it keeps failing


def main():
    srv = ThreadingHTTPServer((HOST, PORT), Handler)
    url = f"http://{HOST}:{PORT}/?t={TOKEN}"
    print(f"Perch running on {HOST}:{PORT} (token-protected)")
    print("  " + url)
    _sd_notify("READY=1")
    threading.Thread(target=_watchdog_loop, daemon=True).start()
    try:
        srv.serve_forever()
    except KeyboardInterrupt:
        pass
